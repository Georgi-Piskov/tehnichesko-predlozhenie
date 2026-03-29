"""
Извличане на съдържанието от трите файла в GRAFIK DG REMONT OP2.
"""
import os
import sys
sys.stdout = open(os.devnull, 'w')  # suppress noise

from docx import Document

BASE = os.path.join(os.path.dirname(__file__), '..', 'GRAFIK DG REMONT OP2')
OUTPUT = os.path.join(BASE, 'docs_output.txt')

lines = []

# 1. КСС (XLS) — use openpyxl if xlsx, or xlrd for xls
lines.append("=" * 80)
lines.append("КСС СМР - база Здравец (1).xls")
lines.append("=" * 80)

try:
    import xlrd
    wb = xlrd.open_workbook(os.path.join(BASE, 'КСС СМР - база Здравец (1).xls'))
    for sheet_idx in range(wb.nsheets):
        sh = wb.sheet_by_index(sheet_idx)
        lines.append(f"\nSheet: {sh.name} ({sh.nrows} rows x {sh.ncols} cols)")
        for row in range(sh.nrows):
            cells = []
            for col in range(sh.ncols):
                val = sh.cell_value(row, col)
                if val:
                    cells.append(str(val).strip()[:80])
                else:
                    cells.append('')
            lines.append(' | '.join(cells))
except ImportError:
    lines.append("xlrd not installed, trying openpyxl...")
    try:
        from openpyxl import load_workbook
        wb = load_workbook(os.path.join(BASE, 'КСС СМР - база Здравец (1).xls'))
        for ws in wb.worksheets:
            lines.append(f"\nSheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip()[:80] if c else '' for c in row]
                lines.append(' | '.join(cells))
    except Exception as e:
        lines.append(f"openpyxl error: {e}")
except Exception as e:
    lines.append(f"xlrd error: {e}")

# 2. Указания  
lines.append("\n" + "=" * 80)
lines.append("Указания (6).docx")
lines.append("=" * 80)

doc = Document(os.path.join(BASE, 'Указания (6).docx'))
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        lines.append(f"[{i}] {p.text.strip()[:200]}")

for ti, t in enumerate(doc.tables):
    lines.append(f"\nТаблица {ti}: {len(t.rows)} x {len(t.columns)}")
    for ri, row in enumerate(t.rows):
        cells = [c.text.strip()[:60] for c in row.cells]
        lines.append(f"  R{ri}: {' | '.join(cells)}")

# 3. Техническо предложение (formatted)
lines.append("\n" + "=" * 80)
lines.append("final_clean_file6_formatted (1).docx — ТЕХНИЧЕСКО ПРЕДЛОЖЕНИЕ")
lines.append("=" * 80)

doc2 = Document(os.path.join(BASE, 'final_clean_file6_formatted (1).docx'))
for i, p in enumerate(doc2.paragraphs):
    if p.text.strip():
        style = p.style.name if p.style else 'N/A'
        lines.append(f"[{i}] ({style}) {p.text.strip()[:200]}")

for ti, t in enumerate(doc2.tables):
    lines.append(f"\nТаблица {ti}: {len(t.rows)} x {len(t.columns)}")
    for ri, row in enumerate(t.rows):
        cells = [c.text.strip()[:60] for c in row.cells]
        lines.append(f"  R{ri}: {' | '.join(cells)}")

sys.stdout = sys.__stdout__
with open(OUTPUT, 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))

print(f"Записано: {OUTPUT}")
print(f"Редове: {len(lines)}")
