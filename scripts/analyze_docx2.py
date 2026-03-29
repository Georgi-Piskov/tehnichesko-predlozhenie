import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

doc = Document(r"e:\VISUAL STUDIO\TEHNICHESKO PREDLOJENIE\ТП_2026-03-26_1637.docx")

# Check styles
print("=== DOCUMENT STYLES IN USE ===")
styles_used = set()
for p in doc.paragraphs:
    styles_used.add(p.style.name)
print("Paragraph styles:", sorted(styles_used))

print("\n=== TOTAL PARAGRAPHS:", len(doc.paragraphs))
print("=== TOTAL TABLES:", len(doc.tables))

print("\n=== PARAGRAPHS WITH Попълнете ===")
count_pop = 0
for i, p in enumerate(doc.paragraphs):
    if "Попълнете" in p.text:
        count_pop += 1
        print("{}: {}".format(i, p.text[:250]))
print("Total with Попълнете:", count_pop)

print("\n=== PARAGRAPHS WITH MARKDOWN TABLE MARKERS (|) ===")
md_table_lines = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith("|") or "---|" in t:
        md_table_lines.append(i)
        if len(md_table_lines) <= 30:
            print("{}: {}".format(i, t[:200]))
print("Total markdown table lines:", len(md_table_lines))

print("\n=== MARKDOWN HEADINGS ===")
heading_count = 0
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith("#"):
        heading_count += 1
        if heading_count <= 30:
            print("{}: {}".format(i, t[:200]))
print("Total markdown headings:", heading_count)

print("\n=== BOLD MARKERS ** ===")
bold_count = 0
for i, p in enumerate(doc.paragraphs):
    if "**" in p.text:
        bold_count += 1
print("Paragraphs with ** markers:", bold_count)

print("\n=== ITALIC MARKERS * (not **) ===")
import re
italic_count = 0
for i, p in enumerate(doc.paragraphs):
    # single * not part of **
    cleaned = p.text.replace("**", "")
    if "*" in cleaned:
        italic_count += 1
print("Paragraphs with * markers:", italic_count)

print("\n=== SAMPLE OF FIRST 30 PARAGRAPHS TEXT ===")
for i in range(min(30, len(doc.paragraphs))):
    print("{}: {}".format(i, doc.paragraphs[i].text[:200]))
