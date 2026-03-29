import sys
from docx import Document

doc = Document(r"e:\VISUAL STUDIO\TEHNICHESKO PREDLOJENIE\ТП_2026-03-26_1637.docx")

# Check styles
print("=== DOCUMENT STYLES IN USE ===")
styles_used = set()
for p in doc.paragraphs:
    styles_used.add(p.style.name)
print("Paragraph styles:", sorted(styles_used))

print("\n=== FIRST 100 PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs[:100]):
    runs_info = []
    for r in p.runs:
        font = r.font
        runs_info.append("font={}, size={}, bold={}, italic={}".format(font.name, font.size, font.bold, font.italic))
    text_preview = p.text[:150] if p.text else "(empty)"
    print("{}: style={} | {}".format(i, p.style.name, text_preview))
    if runs_info:
        for ri in runs_info[:3]:
            print("   run: {}".format(ri))

print("\n=== TABLES ===")
print("Number of tables:", len(doc.tables))
for i, t in enumerate(doc.tables[:5]):
    print("Table {}: {} rows x {} cols".format(i, len(t.rows), len(t.columns)))
    for j, row in enumerate(t.rows[:3]):
        cells = [c.text[:40] for c in row.cells]
        print("  Row {}: {}".format(j, cells))

print("\n=== PARAGRAPHS WITH 'Попълнете' ===")
for i, p in enumerate(doc.paragraphs):
    if "Попълнете" in p.text:
        print("{}: {}".format(i, p.text[:200]))

print("\n=== PARAGRAPHS WITH MARKDOWN TABLE MARKERS ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("|") or "---|" in p.text:
        print("{}: {}".format(i, p.text[:200]))

print("\n=== TOTAL PARAGRAPHS:", len(doc.paragraphs))
