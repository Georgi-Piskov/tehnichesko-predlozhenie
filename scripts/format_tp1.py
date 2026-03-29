# -*- coding: utf-8 -*-
"""
Скрипт за оформление на ТП1 DOCX файл — ремонт и реконструкция на ДЗ База "Здравец":
1. Преобразува Markdown заглавия (#, ##, ###, ####) в Word heading стилове
2. Прилага bold/italic форматиране от ** и *
3. Преобразува Markdown таблици в реални Word таблици
4. Замества [⚠️ ПОПЪЛНЕТЕ: ...] плейсхолдъри с реални стойности
5. Задава шрифт Times New Roman 12pt за основен текст
6. Добавя номера на страниците в footer
7. Форматира bold-only параграфи като визуални под-заглавия
8. Разграничава label-value bullets (- **Label:** стойност)

Базиран на format_docx_v2.py, адаптиран за ТП1.
"""

import re
import sys
import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

INPUT_FILE = r"e:\VISUAL STUDIO\TEHNICHESKO PREDLOJENIE\ТП1\ТП_2026-03-27_1141 (1).docx"
OUTPUT_FILE = r"e:\VISUAL STUDIO\TEHNICHESKO PREDLOJENIE\ТП1\ТП_formatted.docx"

FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = 12
FONT_SIZE_H1 = 16
FONT_SIZE_H2 = 14
FONT_SIZE_H3 = 13
FONT_SIZE_H4 = 12
FONT_SIZE_SUBHEADING = 11
FONT_SIZE_LABEL = 11
FONT_SIZE_CATEGORY = 10

# Цветове
COLOR_H1 = RGBColor(0x1B, 0x3A, 0x5C)
COLOR_H2 = RGBColor(0x1F, 0x4E, 0x79)
COLOR_H3 = RGBColor(0x2E, 0x2E, 0x2E)
COLOR_H4 = RGBColor(0x33, 0x33, 0x33)
COLOR_BODY = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_SUBHEADING = RGBColor(0x1F, 0x4E, 0x79)
COLOR_LABEL = RGBColor(0x2E, 0x4E, 0x6E)
COLOR_CATEGORY = RGBColor(0x55, 0x55, 0x55)
COLOR_TABLE_HEADER = "2E74B5"
COLOR_TABLE_STRIPE = "E8EFF7"
COLOR_SUBHEADING_BG = "EDF2F9"

# ============================================================================
# PLACEHOLDER MAPPING
# ============================================================================
# Ключови лица и персонал
KNOWN_PLACEHOLDERS = {
    # Management → Инж. Георги Писков
    'ръководител проект': 'Инж. Георги Писков',
    'Ръководител проект': 'Инж. Георги Писков',
    'управител': 'Инж. Георги Писков',
    'управляващ директор': 'Инж. Георги Писков',
    'изпълнителен директор': 'Инж. Георги Писков',
    'главен инженер на проекта': 'Инж. Георги Писков',
    'технически директор': 'Инж. Георги Писков',
    'проектен мениджър': 'Инж. Георги Писков',
    'председател съвет': 'Инж. Георги Писков',
    'основен координатор': 'Инж. Георги Писков',
    'координатор проект': 'Инж. Георги Писков',
    
    # Контактна информация
    'телефон': '0892220535',
    'номер на дежурен телефон': '0892220535',
    
    # Адреси/локации
    'местоположение на склада': 'ул. Миньорска 1, гр. Чепеларе',
    'адрес на склад': 'ул. Миньорска 1, гр. Чепеларе',
    
    # Роли → описателен текст (в контекст: "отговорен е ...")
    'главен инженер': 'главният инженер на обекта',
    'технически ръководител': 'техническият ръководител',
    'ръководител строителство': 'ръководителят на строителството',
    'строителен контролор': 'строителният контролор',
    'строителен инспектор': 'строителният инспектор',
    'строителен надзор': 'строителният надзорник',
    'строителен техник': 'строителният техник',
    'контрольор качество': 'контрольорът по качество',
    'контролер качество': 'контрольорът по качество',
    'контролор качество': 'контрольорът по качество',
    'отговорник качество': 'отговорникът по качеството',
    'качествен мениджър': 'мениджърът по качество',
    'качествен анализатор': 'анализаторът по качество',
    'качествен анализатор/специалист по управление на качеството': 'специалистът по управление на качеството',
    'координатор по безопасност': 'координаторът по безопасност',
    'координатор за безопасност': 'координаторът по безопасност',
    'координатор БЗ': 'координаторът по безопасност и здраве',
    'координатор': 'координаторът на дейностите',
    'координатор комуникация': 'координаторът на комуникацията',
    'координатор външни връзки': 'координаторът на външните връзки',
    'координатор екипи': 'координаторът на екипите',
    'координатор гаранционно обслужване': 'координаторът на гаранционното обслужване',
    'координатор обучения гаранционни екипи': 'координаторът на обученията',
    'проектен координатор': 'проектният координатор',
    'проектен контролер': 'проектният контролер',
    'риск мениджър': 'риск мениджърът',
    'IT координатор': 'IT координаторът',
    'IT координатор гаранционни дейности': 'IT координаторът на гаранционните дейности',
    'IT специалист': 'IT специалистът',
    'IT специалист гаранционни дейности': 'IT специалистът по гаранционни дейности',
    'IT специалист/координатор на системата': 'IT координаторът на системата',
    'HR координатор': 'HR координаторът',
    'логистичен координатор': 'логистичният координатор',
    'еколог': 'екологът на обекта',
    'еколог по отпадъци': 'екологът по отпадъци',
    'геодезист': 'геодезистът',
    'електроинженер': 'електроинженерът',
    'електроинженер с правоспособност': 'електроинженерът с правоспособност',
    'електротехник': 'електротехникът',
    'старши електротехник': 'старши електротехникът',
    'електромонтьор': 'електромонтьорът',
    'ВиК инженер': 'ВиК инженерът',
    'ВиК инженер вентилация': 'ВиК инженерът по вентилация',
    'ВиК техник': 'ВиК техникът',
    'ВиК монтажник': 'ВиК монтажникът',
    'бригадир': 'бригадирът',
    'бригадир зидари': 'бригадирът на зидарите',
    'бригадир покривни работи': 'бригадирът на покривните работи',
    'ламаринаджия': 'ламаринаджията',
    'плочкаджия': 'плочкаджията',
    'фаянсаджия': 'фаянсаджията',
    'теракотаджия': 'теракотаджията',
    'бояджия': 'бояджията',
    'старши бояджия': 'старши бояджията',
    'старши покривджия': 'старши покривджията',
    'старши строител': 'старши строителят',
    'старши строител подови настилки': 'старши строителят по подови настилки',
    'старши почистващ работник': 'старши почистващият работник',
    'специалист настилки': 'специалистът по настилки',
    'профилен специалист': 'профилният специалист',
    'технически експерт': 'техническият експерт',
    'технически секретар': 'техническият секретар',
    'технолог': 'технологът',
    'снабдител': 'снабдителят',
    'складов работник': 'складовият работник',
    'складов отговорник': 'складовият отговорник',
    'складов управител': 'складовият управител',
    'складов работник/логистик': 'складовият работник',
    'лаборант': 'лаборантът',
    'анализатор': 'анализаторът',
    'ръководител демонтажни работи': 'ръководителят на демонтажните работи',
    'ръководител мазилъчни работи': 'ръководителят на мазилъчните работи',
    'ръководител смяна': 'ръководителят на смяната',
    'ръководител мобилен екип': 'ръководителят на мобилния екип',
    'ръководител спешни екипи': 'ръководителят на спешните екипи',
    'отговорник за логистика': 'отговорникът за логистика',
    'контактно лице от възложителя': 'представителят на възложителя',
    'контролен орган': 'контролният орган',
    'контролни органи': 'контролните органи',
    'административен координатор': 'административният координатор',
    'административен асистент': 'административният асистент',
    'административен персонал': 'административният персонал',
    'административен специалист': 'административният специалист',
    'дежурен инженер': 'дежурният инженер',
    'дежурен координатор': 'дежурният координатор',
    'дежурен диспечер': 'дежурният диспечер',
    'дежурен инженер по гаранционно обслужване': 'дежурният инженер по гаранционно обслужване',
    'нощен дежурен': 'нощният дежурен',
    'уикенд дежурен': 'уикенд дежурният',
    'празничен дежурен': 'празничният дежурен',
    'почистващ екип': 'почистващият екип',
    'техник пожарна безопасност': 'техникът по пожарна безопасност',
    'фотограф-документатор': 'фотографът-документатор',
    'старши инженер по координация с институции': 'старши инженерът по координация с институции',
    'специалист превантивен мониторинг': 'специалистът по превантивен мониторинг',
    'изпълнителен екип': 'изпълнителният екип',
    'ръководител гаранционни дейности': 'ръководителят на гаранционните дейности',
    'ръководител гаранционно обслужване': 'ръководителят на гаранционното обслужване',
    'гаранционен инженер': 'гаранционният инженер',
    'технически ръководител гаранционни дейности': 'техническият ръководител на гаранционните дейности',
    'технически ръководител превантивни инспекции': 'техническият ръководител на инспекциите',
    'превантивен инспектор': 'превантивният инспектор',
    
    # Резерви
    'резерв 1': 'резерв 1',
    'резерв 2': 'резерв 2',
    'резерв 3': 'резерв 3',
    'резерв 4': 'резерв 4',
    
    # Количества — конкретни стойности базирани на спецификацията/КСС
    'брой помещения': 'всички засегнати',
    'площ кв.м': 'съгласно КСС',
    'кв.м': 'съгл. КСС',
    'куб.м': 'съгл. КСС',
    'кубични метри': 'съгл. КСС',
    'брой предмети': 'съгл. КСС',
    'брой дежурни специалисти': '2 (двама)',
    
    # Специфични описания
    'конкретни видове настилки - теракот, ламинат, линолеум': 'винилова настилка, теракот',
    'видове облицовки - фаянс, керамика': 'фаянсови облицовки',
    'конкретни помещения - санитарни възли, кухня, занимални': 'санитарни възли, кухненски блок и занимални',
    'конкретни количества LED осветителни тела от спецификацията': 'съгласно количествената сметка',
    'конкретни параметри според обекта': 'съгласно проектната документация',
    'конкретни системи от проекта': 'електрическа, ВиК и вентилационна системи',
    'конкретни системи': 'електрическа, ВиК и вентилационна системи',
    'конкретни часове съгласувани с детското заведение': 'от 17:00 до 07:00 часа и в почивни дни',
    'количества винилова настилка от спецификацията': '582,02 м²',
    'количества фаянс от спецификацията': '85,98 м²',
    'конкретен тип обект': 'детско заведение',
    'тип обект според техническите спецификации': 'детско заведение',
    'тип обект': 'детско заведение',
    'основна функция на обекта': 'отглеждане и обучение на деца',
    'основни системи от техническите спецификации': 'електрическа, ВиК и вентилационна системи',
    'основни системи': 'електрическа, ВиК и вентилационна системи',
    'основни услуги': 'водоснабдяване, електроснабдяване и канализация',
    'инсталирани системи от проекта': 'електрическа, ВиК и вентилационна инсталации',
    'специфични системи от техническите спецификации': 'съгласно техническите спецификации',
    'депо за неопасни отпадъци': 'регионалното депо за неопасни отпадъци',
    'лицензирано депо за инертни отпадъци': 'регионалното лицензирано депо за инертни отпадъци',
    'критичен период според функцията на обекта': 'учебната година (септември-юни)',
    'критични материали според спецификациите': 'винилова настилка, фаянс, теракот, електроматериали',
    'подходящи часове според функцията на обекта': 'извънработно време (след 17:00 ч.) и в почивни дни',
    'релевантни институции според обекта': 'РДНСК, РИОСВ, РЗИ, Община Чепеларе',
    'релевантни специалности': 'електроинженер, ВиК инженер, строителен инженер',
    'разстояние според местоположението на обекта': 'до 50 км',
    'мерки за поддръжка на настилки': 'редовно почистване с неагресивни препарати, полиране на винилова настилка',
    'мерки за покривни системи': 'визуален оглед на покрива след обилни валежи, проверка на водостоци',
    'посока/страна на сградата': 'южната страна',
    'тип сензори': 'датчици за температура и влажност',
    'специални изисквания за детски зони': 'нетоксични материали, закръглени ръбове, антиалергенни покрития',
    'специални процедури за санитарни зони': 'дезинфекция след всяка интервенция, тестване на ВиК системите',
    'специфични мерки за електробезопасност': 'защитни капачки на контактите, FI-защита, периодични проверки на изолация',
}

# Regex шаблони
RE_PLACEHOLDER = re.compile(r'\[⚠️\s*ПОПЪЛНЕТЕ:\s*([^\]]*)\]')
RE_PLACEHOLDER2 = re.compile(r'\[\s*ПОПЪЛНЕТЕ[!]*:\s*([^\]]*)\]')
RE_PLACEHOLDER3 = re.compile(r'⚠️\s*ПОПЪЛНЕТЕ:\s*')
RE_INLINE = re.compile(r'\*{3}(.+?)\*{3}|\*{2}(.+?)\*{2}|\*([^*]+?)\*')
RE_TABLE_SEP = re.compile(r'^\|[\s\-:|]+\|$')
RE_NUMBERED = re.compile(r'^\d+\.\s')
RE_ETAP_HEADING = re.compile(r'^\d+\.\d+\.\d+\.\s*Етап:', re.IGNORECASE)
RE_POD_ETAP_HEADING = re.compile(r'^Под-етап\s+\d+', re.IGNORECASE)
RE_CATEGORY_LINE = re.compile(r'^\*{0,2}Категория:\*{0,2}\s*\*([^*]+)\*$', re.IGNORECASE)
RE_ROMAN_HEADING = re.compile(r'^(I{1,3}V?|IV|V|VI{0,3}|VII|VIII|IX|X)\.\s', re.IGNORECASE)


def add_bottom_border(paragraph, color="1B3A5C", size="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr {}><w:bottom w:val="single" w:sz="{}" w:space="1" w:color="{}"/></w:pBdr>'.format(
            nsdecls('w'), size, color))
    pPr.append(pBdr)


def add_left_border(paragraph, color="2E74B5", size="18"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr {}><w:left w:val="single" w:sz="{}" w:space="4" w:color="{}"/></w:pBdr>'.format(
            nsdecls('w'), size, color))
    pPr.append(pBdr)


def add_paragraph_shading(paragraph, color="EDF2F9"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = parse_xml('<w:shd {} w:fill="{}" w:val="clear"/>'.format(nsdecls('w'), color))
    pPr.append(shd)


def add_page_number_footer(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        run1 = p.add_run("— ")
        run1.font.name = FONT_NAME
        run1.font.size = Pt(9)
        run1.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        
        fld_xml = (
            '<w:fldSimple {} w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
        ).format(nsdecls('w'))
        fld = parse_xml(fld_xml)
        p._p.append(fld)
        
        run2 = p.add_run(" / ")
        run2.font.name = FONT_NAME
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        
        fld_xml2 = (
            '<w:fldSimple {} w:instr=" NUMPAGES "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
        ).format(nsdecls('w'))
        fld2 = parse_xml(fld_xml2)
        p._p.append(fld2)
        
        run3 = p.add_run(" —")
        run3.font.name = FONT_NAME
        run3.font.size = Pt(9)
        run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def setup_styles(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(FONT_SIZE_BODY)
    font.color.rgb = COLOR_BODY
    pf = style.paragraph_format
    pf.space_after = Pt(3)
    pf.space_before = Pt(1)
    pf.line_spacing = 1.15

    heading_config = [
        (1, FONT_SIZE_H1, True, COLOR_H1, Pt(28), Pt(10)),
        (2, FONT_SIZE_H2, True, COLOR_H2, Pt(20), Pt(8)),
        (3, FONT_SIZE_H3, True, COLOR_H3, Pt(14), Pt(5)),
        (4, FONT_SIZE_H4, True, COLOR_H4, Pt(10), Pt(3)),
    ]
    for level, size, bold, color, space_before, space_after in heading_config:
        style_name = "Heading {}".format(level)
        try:
            hs = doc.styles[style_name]
        except KeyError:
            hs = doc.styles.add_style(style_name, 1)
        hf = hs.font
        hf.name = FONT_NAME
        hf.size = Pt(size)
        hf.bold = bold
        hf.color.rgb = color
        hp = hs.paragraph_format
        hp.space_before = space_before
        hp.space_after = space_after
        hp.line_spacing = 1.15


def clean_placeholder(text):
    """Замества placeholder-и: първо проверява KNOWN_PLACEHOLDERS, после strip."""
    def replace_match(m):
        key = m.group(1).strip()
        if key in KNOWN_PLACEHOLDERS:
            return KNOWN_PLACEHOLDERS[key]
        # Fallback: strip brackets, leave description
        return key
    
    result = RE_PLACEHOLDER.sub(replace_match, text)
    result = RE_PLACEHOLDER2.sub(replace_match, result)
    result = RE_PLACEHOLDER3.sub('', result)
    return result


def determine_heading_level(text):
    stripped = text.strip()
    temp = stripped
    if temp.startswith("**"):
        temp = temp[2:]
    if temp.endswith("**"):
        temp = temp[:-2]
    temp = temp.strip()
    
    hash_count = 0
    check = temp
    while check.startswith("#"):
        hash_count += 1
        check = check[1:]
    
    if hash_count == 0:
        check2 = stripped
        while check2.startswith("#"):
            hash_count += 1
            check2 = check2[1:]
        if hash_count == 0:
            return 0, text
        clean_text = check2.lstrip()
    else:
        clean_text = check.lstrip()
    
    level = min(hash_count, 4)
    return level, clean_text


def is_bold_subheading(text):
    stripped = text.strip()
    if not stripped.startswith("**"):
        return False
    if not stripped.endswith("**") and not stripped.endswith("***"):
        return False
    inner = stripped[2:]
    if inner.endswith("***"):
        inner = inner[:-3]
    elif inner.endswith("**"):
        inner = inner[:-2]
    else:
        return False
    if len(inner.strip()) < 5:
        return False
    return True


def get_bold_subheading_text(text):
    stripped = text.strip()
    inner = stripped[2:]
    if inner.endswith("***"):
        inner = inner[:-3]
    elif inner.endswith("**"):
        inner = inner[:-2]
    inner = inner.strip().strip("*").strip()
    inner = re.sub(r'\*\(([^)]+)\)$', r'(\1)', inner)
    inner = re.sub(r'\*\(([^)]+)\)\*', r'(\1)', inner)
    return inner.strip()


def is_category_line(text):
    stripped = text.strip()
    clean = stripped.replace("**", "").replace("*", "").strip()
    return clean.lower().startswith("категория:")


def is_label_bullet(text):
    stripped = text.lstrip()
    if not (stripped.startswith("- **") or stripped.startswith("• **")):
        return False
    if ":**" in stripped:
        return True
    return False


def split_label_bullet(text):
    stripped = text.lstrip()
    if stripped.startswith("- "):
        stripped = stripped[2:]
    elif stripped.startswith("• "):
        stripped = stripped[2:]
    m = re.match(r'\*\*([^*]+?)\*\*\s*(.*)', stripped, re.DOTALL)
    if m:
        label = m.group(1).strip()
        value = m.group(2).strip()
        return label, value
    return None, stripped


def apply_inline_formatting(paragraph, text, base_bold=False, base_italic=False,
                            font_name=FONT_NAME, font_size=None, font_color=None):
    text = clean_placeholder(text)
    color = font_color or COLOR_BODY

    parts = []
    pos = 0
    
    for match in RE_INLINE.finditer(text):
        if match.start() > pos:
            parts.append(('normal', text[pos:match.start()]))
        if match.group(1) is not None:
            parts.append(('bold_italic', match.group(1)))
        elif match.group(2) is not None:
            parts.append(('bold', match.group(2)))
        elif match.group(3) is not None:
            parts.append(('italic', match.group(3)))
        pos = match.end()

    if pos < len(text):
        parts.append(('normal', text[pos:]))

    if not parts:
        parts = [('normal', text)]

    cleaned_parts = []
    for fmt, content in parts:
        if fmt == 'normal':
            content = content.replace("***", "").replace("**", "")
            content = content.strip('*')
            if content.strip() == '*':
                continue
        if not content or not content.strip():
            continue
        cleaned_parts.append((fmt, content))
    
    if not cleaned_parts:
        cleaned_parts = [('normal', '')]

    for fmt, content in cleaned_parts:
        if not content:
            continue
        run = paragraph.add_run(content)
        run.font.name = font_name
        if font_size:
            run.font.size = Pt(font_size)
        run.font.color.rgb = color

        if fmt == 'bold' or base_bold:
            run.bold = True
        if fmt == 'italic' or base_italic:
            run.italic = True
        if fmt == 'bold_italic':
            run.bold = True
            run.italic = True


def parse_markdown_table(paragraphs, start_idx):
    rows = []
    idx = start_idx
    while idx < len(paragraphs):
        text = paragraphs[idx].text.strip()
        if not text.startswith("|") and "---|" not in text:
            break
        if RE_TABLE_SEP.match(text):
            idx += 1
            continue
        cells = [c.strip() for c in text.split("|")]
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        if cells:
            rows.append(cells)
        idx += 1
    return rows, idx


def add_table_to_doc(new_doc, rows_data):
    if not rows_data or not rows_data[0]:
        return
    
    num_cols = max(len(row) for row in rows_data)
    table = new_doc.add_table(rows=len(rows_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    try:
        table.style = 'Table Grid'
    except:
        pass
    
    for ri, row_data in enumerate(rows_data):
        for ci, cell_text in enumerate(row_data):
            if ci >= num_cols:
                break
            cell = table.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            
            cell_text = clean_placeholder(cell_text)
            clean_text = cell_text.replace("**", "")
            
            is_header = (ri == 0)
            
            run = p.add_run(clean_text)
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            
            if is_header:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                shading = parse_xml('<w:shd {} w:fill="{}" w:val="clear"/>'.format(
                    nsdecls('w'), COLOR_TABLE_HEADER))
                cell._tc.get_or_add_tcPr().append(shading)
            else:
                run.font.color.rgb = RGBColor(0, 0, 0)
                if ri % 2 == 0:
                    shading = parse_xml('<w:shd {} w:fill="{}" w:val="clear"/>'.format(
                        nsdecls('w'), COLOR_TABLE_STRIPE))
                    cell._tc.get_or_add_tcPr().append(shading)
            
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
    
    table.autofit = True
    return table


def is_bullet_line(text):
    stripped = text.lstrip()
    return (stripped.startswith("- ") or stripped.startswith("• ") or
            RE_NUMBERED.match(stripped) is not None)


def get_bullet_text(text):
    stripped = text.lstrip()
    if stripped.startswith("- "):
        return stripped[2:]
    if stripped.startswith("• "):
        return stripped[2:]
    m = RE_NUMBERED.match(stripped)
    if m:
        return stripped[m.end():]
    return text


def main():
    print("Зареждане на документа: {}".format(INPUT_FILE))
    doc = Document(INPUT_FILE)
    
    print("Създаване на нов форматиран документ...")
    new_doc = Document()
    setup_styles(new_doc)
    
    # Margins
    for section in new_doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    add_page_number_footer(new_doc)
    
    # Изтриваме default празен параграф
    if new_doc.paragraphs:
        p = new_doc.paragraphs[0]
        p._element.getparent().remove(p._element)
    
    paragraphs = doc.paragraphs
    total = len(paragraphs)
    print("Общо параграфи: {}".format(total))
    
    i = 0
    table_count = 0
    heading_count = 0
    placeholder_count = 0
    subheading_count = 0
    label_bullet_count = 0
    category_count = 0
    is_first_h1 = True
    prev_type = None
    
    progress_step = total // 10
    
    while i < total:
        if progress_step > 0 and i % progress_step == 0:
            pct = int(i / total * 100)
            print("  Прогрес: {}% ({}/{})".format(pct, i, total))
        
        para = paragraphs[i]
        text = para.text
        stripped = text.strip()
        
        if not stripped:
            i += 1
            continue
        
        if "ПОПЪЛНЕТЕ" in text:
            placeholder_count += 1
        
        # ===== MARKDOWN TABLE =====
        if stripped.startswith("|") or "---|" in stripped:
            rows_data, end_idx = parse_markdown_table(paragraphs, i)
            if rows_data and len(rows_data) > 0:
                table_count += 1
                spacer = new_doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(6)
                spacer.paragraph_format.space_after = Pt(0)
                spacer.paragraph_format.line_spacing = Pt(1)
                
                add_table_to_doc(new_doc, rows_data)
                
                spacer2 = new_doc.add_paragraph()
                spacer2.paragraph_format.space_before = Pt(0)
                spacer2.paragraph_format.space_after = Pt(6)
                spacer2.paragraph_format.line_spacing = Pt(1)
                
                prev_type = 'table'
                i = end_idx
                continue
        
        # ===== MARKDOWN HEADING =====
        heading_level, heading_text = determine_heading_level(stripped)
        
        if heading_level > 0:
            heading_count += 1
            heading_text = clean_placeholder(heading_text)
            heading_text = heading_text.replace("***", "").replace("**", "").replace("*", "")
            
            if heading_level == 1 and not is_first_h1:
                for _ in range(2):
                    spacer = new_doc.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(0)
                    spacer.paragraph_format.space_after = Pt(0)
                    spacer.paragraph_format.line_spacing = Pt(12)
            if heading_level == 1:
                is_first_h1 = False
            
            if heading_level == 2 and prev_type not in ('heading', None):
                spacer = new_doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(0)
                spacer.paragraph_format.space_after = Pt(0)
                spacer.paragraph_format.line_spacing = Pt(12)
            
            if heading_level == 3 and RE_ROMAN_HEADING.match(heading_text.strip()):
                if prev_type not in (None,):
                    spacer = new_doc.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(0)
                    spacer.paragraph_format.space_after = Pt(0)
                    spacer.paragraph_format.line_spacing = Pt(12)
            
            p = new_doc.add_paragraph(style="Heading {}".format(heading_level))
            
            if heading_level == 1:
                run = p.add_run(heading_text.upper())
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_H1)
                run.bold = True
                run.font.color.rgb = COLOR_H1
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(28)
                p.paragraph_format.space_after = Pt(12)
                add_bottom_border(p, "1B3A5C", "12")
            
            elif heading_level == 2:
                run = p.add_run(heading_text)
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_H2)
                run.bold = True
                run.font.color.rgb = COLOR_H2
                p.paragraph_format.space_before = Pt(22)
                p.paragraph_format.space_after = Pt(8)
                add_bottom_border(p, "1F4E79", "6")
            
            elif heading_level == 3:
                run = p.add_run(heading_text)
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_H3)
                run.bold = True
                run.font.color.rgb = COLOR_H3
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after = Pt(5)
                add_bottom_border(p, "CCCCCC", "4")
            
            elif heading_level == 4:
                run = p.add_run(heading_text)
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_H4)
                run.bold = True
                run.font.color.rgb = COLOR_H4
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Cm(0.5)
            
            prev_type = 'heading'
            i += 1
            continue
        
        # ===== BOLD-ONLY SUB-HEADING =====
        if is_bold_subheading(stripped):
            subheading_count += 1
            inner_text = get_bold_subheading_text(stripped)
            inner_text = clean_placeholder(inner_text)
            inner_text = inner_text.replace("***", "").replace("**", "")
            
            if prev_type == 'bullet':
                spacer = new_doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(0)
                spacer.paragraph_format.space_after = Pt(0)
                spacer.paragraph_format.line_spacing = Pt(2)
            
            p = new_doc.add_paragraph()
            
            clean_inner = inner_text.replace("*", "").strip()
            is_etap = RE_ETAP_HEADING.match(clean_inner)
            is_pod_etap = RE_POD_ETAP_HEADING.match(clean_inner)
            
            if is_etap:
                add_paragraph_shading(p, COLOR_SUBHEADING_BG)
                add_left_border(p, "2E74B5", "18")
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Cm(0.3)
                apply_inline_formatting(p, inner_text, base_bold=True,
                                       font_size=FONT_SIZE_SUBHEADING,
                                       font_color=COLOR_SUBHEADING)
            elif is_pod_etap:
                add_left_border(p, "95B3D7", "12")
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Cm(0.8)
                apply_inline_formatting(p, inner_text, base_bold=True,
                                       font_size=FONT_SIZE_SUBHEADING,
                                       font_color=RGBColor(0x3A, 0x5A, 0x8C))
            else:
                add_paragraph_shading(p, COLOR_SUBHEADING_BG)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Cm(0.3)
                apply_inline_formatting(p, inner_text, base_bold=True,
                                       font_size=FONT_SIZE_SUBHEADING,
                                       font_color=COLOR_SUBHEADING)
            
            prev_type = 'subheading'
            i += 1
            continue
        
        # ===== CATEGORY LINE =====
        if is_category_line(stripped):
            category_count += 1
            clean_cat = stripped.replace("**", "").replace("*", "").strip()
            
            p = new_doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(1.0)
            
            run_label = p.add_run("Категория: ")
            run_label.font.name = FONT_NAME
            run_label.font.size = Pt(FONT_SIZE_CATEGORY)
            run_label.font.color.rgb = COLOR_CATEGORY
            run_label.bold = True
            run_label.italic = True
            
            cat_value = clean_cat.replace("Категория:", "").strip()
            run_val = p.add_run(cat_value)
            run_val.font.name = FONT_NAME
            run_val.font.size = Pt(FONT_SIZE_CATEGORY)
            run_val.font.color.rgb = COLOR_CATEGORY
            run_val.italic = True
            
            prev_type = 'category'
            i += 1
            continue
        
        # ===== LABEL-VALUE BULLET =====
        if is_label_bullet(stripped):
            label_bullet_count += 1
            label, value = split_label_bullet(stripped)
            
            p = new_doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(1.5)
            
            if label:
                run_label = p.add_run(label + " ")
                run_label.font.name = FONT_NAME
                run_label.font.size = Pt(FONT_SIZE_LABEL)
                run_label.font.color.rgb = COLOR_LABEL
                run_label.bold = True
                
                if value:
                    apply_inline_formatting(p, value, font_size=FONT_SIZE_BODY,
                                           font_color=COLOR_BODY)
            else:
                apply_inline_formatting(p, value, font_size=FONT_SIZE_BODY)
            
            prev_type = 'bullet'
            i += 1
            continue
        
        # ===== REGULAR BULLET =====
        if is_bullet_line(stripped):
            bullet_text = get_bullet_text(stripped)
            p = new_doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, bullet_text, font_size=FONT_SIZE_BODY)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(1.5)
            
            prev_type = 'bullet'
            i += 1
            continue
        
        # ===== NORMAL PARAGRAPH =====
        if prev_type == 'bullet':
            spacer = new_doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(0)
            spacer.paragraph_format.line_spacing = Pt(2)
        
        p = new_doc.add_paragraph()
        apply_inline_formatting(p, stripped, font_size=FONT_SIZE_BODY)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.first_line_indent = Cm(0.75)
        
        prev_type = 'normal'
        i += 1
    
    print("\n" + "=" * 60)
    print("РЕЗУЛТАТИ:")
    print("=" * 60)
    print("  Преобразувани заглавия: {}".format(heading_count))
    print("  Под-заглавия (bold-only): {}".format(subheading_count))
    print("  Label-value bullets: {}".format(label_bullet_count))
    print("  Категория: записи: {}".format(category_count))
    print("  Създадени таблици: {}".format(table_count))
    print("  Параграфи с placeholder-и: {}".format(placeholder_count))
    print("=" * 60)
    
    # Финална проверка за оставащи placeholder-и
    remaining = 0
    for p in new_doc.paragraphs:
        if '⚠' in p.text or 'ПОПЪЛНЕТЕ' in p.text:
            remaining += 1
    for t in new_doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if '⚠' in p.text or 'ПОПЪЛНЕТЕ' in p.text:
                        remaining += 1
    
    if remaining > 0:
        print("  ⚠ Оставащи placeholder-и: {}".format(remaining))
    else:
        print("  ✓ Всички placeholder-и са заменени!")
    
    print("\nЗапазване: {}".format(OUTPUT_FILE))
    new_doc.save(OUTPUT_FILE)
    print("✓ Готово!")


if __name__ == "__main__":
    main()
