# -*- coding: utf-8 -*-
"""
Скрипт за оформление на генериран DOCX файл - v2 (професионално форматиране):
1. Преобразува Markdown заглавия (#, ##, ###, ####) в Word heading стилове
2. Прилага bold/italic форматиране от ** и *
3. Преобразува Markdown таблици в реални Word таблици
4. Премахва [⚠️ ПОПЪЛНЕТЕ: ...] плейсхолдъри, оставяйки само длъжността
5. Задава шрифт Times New Roman 12pt за основен текст
6. Добавя page breaks преди основни секции
7. Добавя номера на страниците в footer
8. Форматира bold-only параграфи като визуални под-заглавия
9. Разграничава label-value bullets (- **Label:** стойност)
10. Форматира "Категория:" редове като metadata
11. Добавя по-голямо разстояние между логически блокове
"""

import re
import sys
import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

INPUT_FILE = r"e:\VISUAL STUDIO\TEHNICHESKO PREDLOJENIE\ТП_2026-03-26_1637.docx"
OUTPUT_FILE = r"e:\VISUAL STUDIO\TEHNICHESKO PREDLOJENIE\ТП_2026-03-26_1637_formatted_v2.docx"

FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = 12
FONT_SIZE_H1 = 16
FONT_SIZE_H2 = 14
FONT_SIZE_H3 = 13
FONT_SIZE_H4 = 12
FONT_SIZE_SUBHEADING = 11   # Bold-only sub-heading lines
FONT_SIZE_LABEL = 11        # Label text in bullets
FONT_SIZE_CATEGORY = 10     # Category metadata

# Цветове
COLOR_H1 = RGBColor(0x1B, 0x3A, 0x5C)      # тъмносин
COLOR_H2 = RGBColor(0x1F, 0x4E, 0x79)      # среден син
COLOR_H3 = RGBColor(0x2E, 0x2E, 0x2E)      # тъмносив
COLOR_H4 = RGBColor(0x33, 0x33, 0x33)      # сив
COLOR_BODY = RGBColor(0x1A, 0x1A, 0x1A)    # почти черно
COLOR_SUBHEADING = RGBColor(0x1F, 0x4E, 0x79)  # син за под-заглавия
COLOR_LABEL = RGBColor(0x2E, 0x4E, 0x6E)   # тъмносин за label-и в bullets
COLOR_CATEGORY = RGBColor(0x55, 0x55, 0x55) # сив за Категория:
COLOR_TABLE_HEADER = "2E74B5"               # Син за таблични заглавия
COLOR_TABLE_STRIPE = "E8EFF7"               # Бледосин за zebra stripe
COLOR_SUBHEADING_BG = "EDF2F9"              # Фон за под-заглавия

# Шаблони
RE_PLACEHOLDER = re.compile(r'\[⚠️\s*ПОПЪЛНЕТЕ:\s*([^\]]*)\]')
RE_PLACEHOLDER2 = re.compile(r'\[\s*ПОПЪЛНЕТЕ[!]*:\s*([^\]]*)\]')
RE_PLACEHOLDER3 = re.compile(r'⚠️\s*ПОПЪЛНЕТЕ:\s*')
RE_INLINE = re.compile(r'\*{3}(.+?)\*{3}|\*{2}(.+?)\*{2}|\*([^*]+?)\*')
RE_TABLE_SEP = re.compile(r'^\|[\s\-:|]+\|$')
RE_NUMBERED = re.compile(r'^\d+\.\s')
RE_ETAP_HEADING = re.compile(r'^\d+\.\d+\.\d+\.\s*Етап:', re.IGNORECASE)
RE_POD_ETAP_HEADING = re.compile(r'^Под-етап\s+\d+', re.IGNORECASE)
RE_CATEGORY_LINE = re.compile(r'^\*{0,2}Категория:\*{0,2}\s*\*([^*]+)\*$', re.IGNORECASE)
RE_ROMAN_HEADING = re.compile(r'^(I{1,3}V?|IV|V|VI{0,3}|VII|VIII|IX|X)\.\s', re.IGNORECASE)


def add_bottom_border(paragraph, color="1B3A5C", size="8"):
    """Добавя долна линия под параграф."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr {}><w:bottom w:val="single" w:sz="{}" w:space="1" w:color="{}"/></w:pBdr>'.format(
            nsdecls('w'), size, color))
    pPr.append(pBdr)


def add_left_border(paragraph, color="2E74B5", size="18"):
    """Добавя лява линия (left border/indent bar) на параграф."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr {}><w:left w:val="single" w:sz="{}" w:space="4" w:color="{}"/></w:pBdr>'.format(
            nsdecls('w'), size, color))
    pPr.append(pBdr)


def add_paragraph_shading(paragraph, color="EDF2F9"):
    """Добавя фоново оцветяване на параграф."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = parse_xml('<w:shd {} w:fill="{}" w:val="clear"/>'.format(nsdecls('w'), color))
    pPr.append(shd)


def add_page_number_footer(doc):
    """Добавя номерация на страниците в footer-а."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        # Добавяме "Страница " text
        run1 = p.add_run("— ")
        run1.font.name = FONT_NAME
        run1.font.size = Pt(9)
        run1.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        
        # PAGE field
        fld_xml = (
            '<w:fldSimple {} w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
        ).format(nsdecls('w'))
        fld = parse_xml(fld_xml)
        p._p.append(fld)
        
        # " от "
        run2 = p.add_run(" / ")
        run2.font.name = FONT_NAME
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        
        # NUMPAGES field
        fld_xml2 = (
            '<w:fldSimple {} w:instr=" NUMPAGES "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
        ).format(nsdecls('w'))
        fld2 = parse_xml(fld_xml2)
        p._p.append(fld2)
        
        run3 = p.add_run(" —")
        run3.font.name = FONT_NAME
        run3.font.size = Pt(9)
        run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def setup_styles(doc):
    """Настройва стиловете на документа."""
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(FONT_SIZE_BODY)
    font.color.rgb = COLOR_BODY
    pf = style.paragraph_format
    pf.space_after = Pt(3)
    pf.space_before = Pt(1)
    pf.line_spacing = 1.15

    heading_config = [
        (1, FONT_SIZE_H1, True, COLOR_H1, Pt(28), Pt(10)),
        (2, FONT_SIZE_H2, True, COLOR_H2, Pt(20), Pt(8)),
        (3, FONT_SIZE_H3, True, COLOR_H3, Pt(14), Pt(5)),
        (4, FONT_SIZE_H4, True, COLOR_H4, Pt(10), Pt(3)),
    ]
    for level, size, bold, color, space_before, space_after in heading_config:
        style_name = "Heading {}".format(level)
        try:
            hs = doc.styles[style_name]
        except KeyError:
            hs = doc.styles.add_style(style_name, 1)
        hf = hs.font
        hf.name = FONT_NAME
        hf.size = Pt(size)
        hf.bold = bold
        hf.color.rgb = color
        hp = hs.paragraph_format
        hp.space_before = space_before
        hp.space_after = space_after
        hp.line_spacing = 1.15


def clean_placeholder(text):
    """Премахва ПОПЪЛНЕТЕ placeholder-и, оставяйки само длъжността."""
    result = RE_PLACEHOLDER.sub(lambda m: m.group(1).strip(), text)
    result = RE_PLACEHOLDER2.sub(lambda m: m.group(1).strip(), result)
    result = RE_PLACEHOLDER3.sub('', result)
    return result


def determine_heading_level(text):
    """Определя нивото на heading от markdown синтаксис."""
    stripped = text.strip()
    temp = stripped
    if temp.startswith("**"):
        temp = temp[2:]
    if temp.endswith("**"):
        temp = temp[:-2]
    temp = temp.strip()
    
    hash_count = 0
    check = temp
    while check.startswith("#"):
        hash_count += 1
        check = check[1:]
    
    if hash_count == 0:
        check2 = stripped
        while check2.startswith("#"):
            hash_count += 1
            check2 = check2[1:]
        if hash_count == 0:
            return 0, text
        clean_text = check2.lstrip()
    else:
        clean_text = check.lstrip()
    
    level = min(hash_count, 4)
    return level, clean_text


def is_bold_subheading(text):
    """
    Проверява дали текстът е bold-only параграф - потенциално под-заглавие.
    Формати: **X.X.X. Етап: текст *(категория)*** или **Под-етап X.X.X.X: ...**
    """
    stripped = text.strip()
    if not stripped.startswith("**"):
        return False
    if not stripped.endswith("**"):
        # Може да завършва с ***
        if not stripped.endswith("***"):
            return False
    # Проверяваме дали целият текст е обвит в **...**
    inner = stripped[2:]
    if inner.endswith("***"):
        inner = inner[:-3]
    elif inner.endswith("**"):
        inner = inner[:-2]
    else:
        return False
    # Трябва да има поне 5 символа съдържание
    if len(inner.strip()) < 5:
        return False
    return True


def get_bold_subheading_text(text):
    """Извлича текста от bold-only под-заглавие."""
    stripped = text.strip()
    inner = stripped[2:]  # remove leading **
    if inner.endswith("***"):
        inner = inner[:-3]
    elif inner.endswith("**"):
        inner = inner[:-2]
    # Почистваме стряй * от началото/края
    inner = inner.strip().strip("*").strip()
    # Поправяме непълни italic маркери: *(text) → (text)
    inner = re.sub(r'\*\(([^)]+)\)$', r'(\1)', inner)
    inner = re.sub(r'\*\(([^)]+)\)\*', r'(\1)', inner)
    return inner.strip()


def is_category_line(text):
    """Проверява дали е ред от вида **Категория:** *инициране*"""
    stripped = text.strip()
    clean = stripped.replace("**", "").replace("*", "").strip()
    return clean.lower().startswith("категория:")


def is_label_bullet(text):
    """
    Проверява дали е label-value bullet: - **Label:** стойност
    """
    stripped = text.lstrip()
    if not (stripped.startswith("- **") or stripped.startswith("• **")):
        return False
    if ":**" in stripped:
        return True
    return False


def split_label_bullet(text):
    """
    Разделя label-value bullet на (label, value).
    Входни примери: '- **Дейност:** описание текст'
    Връща: ('Дейност:', 'описание текст')
    """
    stripped = text.lstrip()
    # Премахваме водещи - или •
    if stripped.startswith("- "):
        stripped = stripped[2:]
    elif stripped.startswith("• "):
        stripped = stripped[2:]
    
    # Намираме **Label:** pattern
    m = re.match(r'\*\*([^*]+?)\*\*\s*(.*)', stripped, re.DOTALL)
    if m:
        label = m.group(1).strip()
        value = m.group(2).strip()
        return label, value
    return None, stripped


def apply_inline_formatting(paragraph, text, base_bold=False, base_italic=False, 
                            font_name=FONT_NAME, font_size=None, font_color=None):
    """Парсва markdown inline форматиране и добавя runs."""
    text = clean_placeholder(text)
    color = font_color or COLOR_BODY

    parts = []
    pos = 0
    
    for match in RE_INLINE.finditer(text):
        if match.start() > pos:
            parts.append(('normal', text[pos:match.start()]))
        if match.group(1) is not None:
            parts.append(('bold_italic', match.group(1)))
        elif match.group(2) is not None:
            parts.append(('bold', match.group(2)))
        elif match.group(3) is not None:
            parts.append(('italic', match.group(3)))
        pos = match.end()

    if pos < len(text):
        parts.append(('normal', text[pos:]))

    if not parts:
        parts = [('normal', text)]

    cleaned_parts = []
    for fmt, content in parts:
        if fmt == 'normal':
            content = content.replace("***", "").replace("**", "")
            content = content.strip('*')
            if content.strip() == '*':
                continue
        if not content or not content.strip():
            continue
        cleaned_parts.append((fmt, content))
    
    if not cleaned_parts:
        cleaned_parts = [('normal', '')]

    for fmt, content in cleaned_parts:
        if not content:
            continue
        run = paragraph.add_run(content)
        run.font.name = font_name
        if font_size:
            run.font.size = Pt(font_size)
        run.font.color.rgb = color

        if fmt == 'bold' or base_bold:
            run.bold = True
        if fmt == 'italic' or base_italic:
            run.italic = True
        if fmt == 'bold_italic':
            run.bold = True
            run.italic = True


def parse_markdown_table(paragraphs, start_idx):
    """Парсва markdown таблица."""
    rows = []
    idx = start_idx
    while idx < len(paragraphs):
        text = paragraphs[idx].text.strip()
        if not text.startswith("|") and "---|" not in text:
            break
        if RE_TABLE_SEP.match(text):
            idx += 1
            continue
        cells = [c.strip() for c in text.split("|")]
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        if cells:
            rows.append(cells)
        idx += 1
    return rows, idx


def add_table_to_doc(new_doc, rows_data):
    """Създава форматирана Word таблица."""
    if not rows_data or not rows_data[0]:
        return
    
    num_cols = max(len(row) for row in rows_data)
    table = new_doc.add_table(rows=len(rows_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    try:
        table.style = 'Table Grid'
    except:
        pass
    
    for ri, row_data in enumerate(rows_data):
        for ci, cell_text in enumerate(row_data):
            if ci >= num_cols:
                break
            cell = table.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            
            cell_text = clean_placeholder(cell_text)
            clean_text = cell_text.replace("**", "")
            
            is_header = (ri == 0)
            
            run = p.add_run(clean_text)
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            
            if is_header:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                shading = parse_xml('<w:shd {} w:fill="{}" w:val="clear"/>'.format(
                    nsdecls('w'), COLOR_TABLE_HEADER))
                cell._tc.get_or_add_tcPr().append(shading)
            else:
                run.font.color.rgb = RGBColor(0, 0, 0)
                if ri % 2 == 0:
                    shading = parse_xml('<w:shd {} w:fill="{}" w:val="clear"/>'.format(
                        nsdecls('w'), COLOR_TABLE_STRIPE))
                    cell._tc.get_or_add_tcPr().append(shading)
            
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
    
    table.autofit = True
    return table


def is_bullet_line(text):
    """Проверява дали е маркиран списък."""
    stripped = text.lstrip()
    return (stripped.startswith("- ") or stripped.startswith("• ") or 
            RE_NUMBERED.match(stripped) is not None)


def get_bullet_text(text):
    """Извлича текста от маркиран списък."""
    stripped = text.lstrip()
    if stripped.startswith("- "):
        return stripped[2:]
    if stripped.startswith("• "):
        return stripped[2:]
    m = RE_NUMBERED.match(stripped)
    if m:
        return stripped[m.end():]
    return text


def add_page_break(new_doc):
    """Добавя page break."""
    p = new_doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type=7)  # WD_BREAK.PAGE
    return p


def main():
    print("Зареждане на документа...")
    doc = Document(INPUT_FILE)
    
    print("Настройване на стилове...")
    new_doc = Document()
    setup_styles(new_doc)
    
    # Margins
    for section in new_doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    # Page numbers
    add_page_number_footer(new_doc)
    
    # Изтриваме default празен параграф
    if new_doc.paragraphs:
        p = new_doc.paragraphs[0]
        p._element.getparent().remove(p._element)
    
    paragraphs = doc.paragraphs
    total = len(paragraphs)
    print("Общо параграфи: {}".format(total))
    
    i = 0
    table_count = 0
    heading_count = 0
    placeholder_count = 0
    subheading_count = 0
    label_bullet_count = 0
    category_count = 0
    page_break_count = 0
    is_first_h1 = True
    prev_type = None  # Track previous paragraph type for spacing
    
    # Import for page break
    from docx.enum.text import WD_BREAK as WdBreak
    
    while i < total:
        para = paragraphs[i]
        text = para.text
        stripped = text.strip()
        
        if not stripped:
            i += 1
            continue
        
        # Placeholder count
        if "ПОПЪЛНЕТЕ" in text:
            placeholder_count += 1
        
        # ===== MARKDOWN TABLE =====
        if stripped.startswith("|") or "---|" in stripped:
            rows_data, end_idx = parse_markdown_table(paragraphs, i)
            if rows_data and len(rows_data) > 0:
                table_count += 1
                spacer = new_doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(6)
                spacer.paragraph_format.space_after = Pt(0)
                sf = spacer.paragraph_format
                sf.line_spacing = Pt(1)
                
                add_table_to_doc(new_doc, rows_data)
                
                spacer2 = new_doc.add_paragraph()
                spacer2.paragraph_format.space_before = Pt(0)
                spacer2.paragraph_format.space_after = Pt(6)
                sf2 = spacer2.paragraph_format
                sf2.line_spacing = Pt(1)
                
                prev_type = 'table'
                i = end_idx
                continue
        
        # ===== MARKDOWN HEADING =====
        heading_level, heading_text = determine_heading_level(stripped)
        
        if heading_level > 0:
            heading_count += 1
            heading_text = clean_placeholder(heading_text)
            heading_text = heading_text.replace("***", "").replace("**", "").replace("*", "")
            
            # Вместо page break — празни редове за визуално разделяне
            if heading_level == 1 and not is_first_h1:
                # Два празни реда преди H1
                for _ in range(2):
                    spacer = new_doc.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(0)
                    spacer.paragraph_format.space_after = Pt(0)
                    spacer.paragraph_format.line_spacing = Pt(12)
            if heading_level == 1:
                is_first_h1 = False
            
            # Празен ред преди H2 (нова основна секция)
            if heading_level == 2 and prev_type not in ('heading', None):
                spacer = new_doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(0)
                spacer.paragraph_format.space_after = Pt(0)
                spacer.paragraph_format.line_spacing = Pt(12)
            
            # Празен ред преди H3 с римски цифри (жизнен цикъл секции)
            if heading_level == 3 and RE_ROMAN_HEADING.match(heading_text.strip()):
                if prev_type not in (None,):
                    spacer = new_doc.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(0)
                    spacer.paragraph_format.space_after = Pt(0)
                    spacer.paragraph_format.line_spacing = Pt(12)
            
            p = new_doc.add_paragraph(style="Heading {}".format(heading_level))
            
            if heading_level == 1:
                run = p.add_run(heading_text.upper())
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_H1)
                run.bold = True
                run.font.color.rgb = COLOR_H1
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(28)
                p.paragraph_format.space_after = Pt(12)
                add_bottom_border(p, "1B3A5C", "12")
            
            elif heading_level == 2:
                run = p.add_run(heading_text)
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_H2)
                run.bold = True
                run.font.color.rgb = COLOR_H2
                p.paragraph_format.space_before = Pt(22)
                p.paragraph_format.space_after = Pt(8)
                add_bottom_border(p, "1F4E79", "6")
            
            elif heading_level == 3:
                run = p.add_run(heading_text)
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_H3)
                run.bold = True
                run.font.color.rgb = COLOR_H3
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after = Pt(5)
                add_bottom_border(p, "CCCCCC", "4")
            
            elif heading_level == 4:
                run = p.add_run(heading_text)
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_H4)
                run.bold = True
                run.font.color.rgb = COLOR_H4
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Cm(0.5)
            
            prev_type = 'heading'
            i += 1
            continue
        
        # ===== BOLD-ONLY SUB-HEADING =====
        # Формати: **X.X.X. Етап: text** или **Под-етап X.X.X.X: text**
        if is_bold_subheading(stripped):
            subheading_count += 1
            inner_text = get_bold_subheading_text(stripped)
            inner_text = clean_placeholder(inner_text)
            # Премахваме вътрешни * маркери
            inner_text = inner_text.replace("***", "").replace("**", "")
            # Оставяме единични * за italic
            
            # Добавяме малко повече разстояние ако предишният не е heading
            if prev_type == 'bullet':
                spacer = new_doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(0)
                spacer.paragraph_format.space_after = Pt(0)
                sp_fmt = spacer.paragraph_format
                sp_fmt.line_spacing = Pt(2)
            
            p = new_doc.add_paragraph()
            
            # Разпознаваме типа sub-heading
            clean_inner = inner_text.replace("*", "").strip()
            is_etap = RE_ETAP_HEADING.match(clean_inner)
            is_pod_etap = RE_POD_ETAP_HEADING.match(clean_inner)
            
            if is_etap:
                # X.X.X. Етап: ... — по-видим, с лява линия и фон
                add_paragraph_shading(p, COLOR_SUBHEADING_BG)
                add_left_border(p, "2E74B5", "18")
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Cm(0.3)
                
                # Парсваме за italic части *(категория)*
                apply_inline_formatting(p, inner_text, base_bold=True,
                                       font_size=FONT_SIZE_SUBHEADING,
                                       font_color=COLOR_SUBHEADING)
            
            elif is_pod_etap:
                # Под-етап — по-малък, с лява линия но без фон  
                add_left_border(p, "95B3D7", "12")
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Cm(0.8)
                
                apply_inline_formatting(p, inner_text, base_bold=True,
                                       font_size=FONT_SIZE_SUBHEADING,
                                       font_color=RGBColor(0x3A, 0x5A, 0x8C))
            
            else:
                # Друг вид bold параграф — generic sub-heading
                add_paragraph_shading(p, COLOR_SUBHEADING_BG)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Cm(0.3)
                
                apply_inline_formatting(p, inner_text, base_bold=True,
                                       font_size=FONT_SIZE_SUBHEADING,
                                       font_color=COLOR_SUBHEADING)
            
            prev_type = 'subheading'
            i += 1
            continue
        
        # ===== CATEGORY LINE =====
        if is_category_line(stripped):
            category_count += 1
            # Извличаме категорията
            clean_cat = stripped.replace("**", "").replace("*", "").strip()
            
            p = new_doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(1.0)
            
            # "Категория:" bold
            run_label = p.add_run("Категория: ")
            run_label.font.name = FONT_NAME
            run_label.font.size = Pt(FONT_SIZE_CATEGORY)
            run_label.font.color.rgb = COLOR_CATEGORY
            run_label.bold = True
            run_label.italic = True
            
            # Стойност — italic
            cat_value = clean_cat.replace("Категория:", "").strip()
            run_val = p.add_run(cat_value)
            run_val.font.name = FONT_NAME
            run_val.font.size = Pt(FONT_SIZE_CATEGORY)
            run_val.font.color.rgb = COLOR_CATEGORY
            run_val.italic = True
            
            prev_type = 'category'
            i += 1
            continue
        
        # ===== LABEL-VALUE BULLET =====
        if is_label_bullet(stripped):
            label_bullet_count += 1
            label, value = split_label_bullet(stripped)
            
            p = new_doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(1.5)
            
            if label:
                # Label — bold, color
                run_label = p.add_run(label + " ")
                run_label.font.name = FONT_NAME
                run_label.font.size = Pt(FONT_SIZE_LABEL)
                run_label.font.color.rgb = COLOR_LABEL
                run_label.bold = True
                
                # Value — normal
                if value:
                    apply_inline_formatting(p, value, font_size=FONT_SIZE_BODY,
                                           font_color=COLOR_BODY)
            else:
                apply_inline_formatting(p, value, font_size=FONT_SIZE_BODY)
            
            prev_type = 'bullet'
            i += 1
            continue
        
        # ===== REGULAR BULLET =====
        if is_bullet_line(stripped):
            bullet_text = get_bullet_text(stripped)
            p = new_doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, bullet_text, font_size=FONT_SIZE_BODY)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(1.5)
            
            prev_type = 'bullet'
            i += 1
            continue
        
        # ===== NORMAL PARAGRAPH =====
        # Добавяме разстояние ако преминаваме от bullets към нормален текст
        if prev_type == 'bullet':
            spacer = new_doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(0)
            sp_fmt = spacer.paragraph_format
            sp_fmt.line_spacing = Pt(2)
        
        p = new_doc.add_paragraph()
        apply_inline_formatting(p, stripped, font_size=FONT_SIZE_BODY)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.first_line_indent = Cm(0.75)
        
        prev_type = 'normal'
        i += 1
    
    print("=" * 50)
    print("Преобразувани заглавия: {}".format(heading_count))
    print("Под-заглавия (bold-only): {}".format(subheading_count))
    print("Label-value bullets: {}".format(label_bullet_count))
    print("Категория: записи: {}".format(category_count))
    print("Създадени таблици: {}".format(table_count))
    print("Коригирани placeholder-и: {}".format(placeholder_count))
    print("Page breaks: {}".format(page_break_count))
    print("=" * 50)
    
    print("Запазване на {}...".format(OUTPUT_FILE))
    new_doc.save(OUTPUT_FILE)
    print("Готово!")


if __name__ == "__main__":
    main()
