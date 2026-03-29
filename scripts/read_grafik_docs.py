import sys
from docx import Document
import PyPDF2

out = open("GRAFIK DG_KUHNQ/docs_output.txt", "w", encoding="utf-8")

# 1. Read price proposal
out.write("=" * 80 + "\n")
out.write("CENOVO PREDLOJENIE OP 2\n")
out.write("=" * 80 + "\n")
doc = Document("GRAFIK DG_KUHNQ/Образец № 3.2. - Ценово предложение ОП 2 – ново.docx")
for i, table in enumerate(doc.tables):
    out.write(f"\n--- Table {i+1} ---\n")
    for row in table.rows:
        cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
        out.write(" | ".join(cells) + "\n")
out.write("\n--- Paragraphs ---\n")
for p in doc.paragraphs:
    t = p.text.strip()
    if t:
        out.write(t + "\n")

# 2. Read technical proposal
out.write("\n" + "=" * 80 + "\n")
out.write("TEHNICHESKO PREDLOJENIE\n")
out.write("=" * 80 + "\n")
doc2 = Document("GRAFIK DG_KUHNQ/Кухня_technical_proposal_final_clean (1).docx")
for p in doc2.paragraphs:
    t = p.text.strip()
    if t:
        out.write(t + "\n")
for i, table in enumerate(doc2.tables):
    out.write(f"\n--- Table {i+1} ---\n")
    for row in table.rows:
        cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
        out.write(" | ".join(cells) + "\n")

# 3. Read PDF instructions
out.write("\n" + "=" * 80 + "\n")
out.write("UKAZANIQ (PDF)\n")
out.write("=" * 80 + "\n")
with open("GRAFIK DG_KUHNQ/Указания ново (1).pdf", "rb") as f:
    reader = PyPDF2.PdfReader(f)
    for page in reader.pages:
        text = page.extract_text()
        if text:
            out.write(text + "\n")

out.close()
print("Done - output in GRAFIK DG_KUHNQ/docs_output.txt")
