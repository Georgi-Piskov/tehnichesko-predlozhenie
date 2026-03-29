"""
Извличане на пълното съдържание на ТП от ТП1 папката,
включително стилове, таблици, и форматиране.
"""
from docx import Document
from docx.shared import Pt, Cm, Emu
import os

INPUT = os.path.join(os.path.dirname(__file__), '..', 'ТП1', 'ТП_2026-03-27_1141 (1).docx')
OUTPUT = os.path.join(os.path.dirname(__file__), '..', 'ТП1', 'tp1_full_analysis.txt')

doc = Document(INPUT)
lines = []

# Document sections (page setup)
lines.append("=" * 80)
lines.append("СТРАНИЦА / СЕКЦИИ")
lines.append("=" * 80)
for i, section in enumerate(doc.sections):
    lines.append(f"Section {i}: width={section.page_width}, height={section.page_height}")
    lines.append(f"  margins: top={section.top_margin}, bottom={section.bottom_margin}, left={section.left_margin}, right={section.right_margin}")

# Styles used
lines.append("\n" + "=" * 80)
lines.append("СТИЛОВЕ В УПОТРЕБА")
lines.append("=" * 80)
style_counts = {}
for p in doc.paragraphs:
    sname = p.style.name if p.style else 'None'
    style_counts[sname] = style_counts.get(sname, 0) + 1
for s, c in sorted(style_counts.items(), key=lambda x: -x[1]):
    lines.append(f"  {s}: {c} пъти")

# Full content with paragraph index, style, alignment, font details
lines.append("\n" + "=" * 80)
lines.append("ПЪЛНО СЪДЪРЖАНИЕ")
lines.append("=" * 80)
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else 'None'
    align = str(p.alignment) if p.alignment else 'None'
    indent_left = p.paragraph_format.left_indent
    indent_first = p.paragraph_format.first_line_indent
    space_before = p.paragraph_format.space_before
    space_after = p.paragraph_format.space_after

    # Run details
    run_info = []
    for r in p.runs:
        font = r.font
        rdetails = []
        if font.bold: rdetails.append("B")
        if font.italic: rdetails.append("I")
        if font.underline: rdetails.append("U")
        if font.size: rdetails.append(f"sz={font.size}")
        if font.name: rdetails.append(f"fn={font.name}")
        if font.color and font.color.rgb: rdetails.append(f"clr={font.color.rgb}")
        run_info.append(f"[{','.join(rdetails) if rdetails else '-'}]'{r.text}'")

    meta = f"[{i}] ({style}|align={align}"
    if indent_left: meta += f"|indL={indent_left}"
    if indent_first: meta += f"|indF={indent_first}"
    meta += ")"

    text = p.text.strip()
    if text or run_info:
        lines.append(meta)
        if run_info:
            lines.append("  RUNS: " + " | ".join(run_info[:10]))  # limit runs shown
            if len(run_info) > 10:
                lines.append(f"  ... +{len(run_info)-10} more runs")
        lines.append(f"  TEXT: {text}")
        lines.append("")

# Tables
lines.append("\n" + "=" * 80)
lines.append("ТАБЛИЦИ")
lines.append("=" * 80)
for ti, table in enumerate(doc.tables):
    lines.append(f"\nТаблица {ti}: {len(table.rows)} реда x {len(table.columns)} колони")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip()[:60] for cell in row.cells]
        lines.append(f"  Row {ri}: {' | '.join(cells)}")

# Placeholders check
lines.append("\n" + "=" * 80)
lines.append("PLACEHOLDER МАРКЕРИ")
lines.append("=" * 80)
for i, p in enumerate(doc.paragraphs):
    if '⚠' in p.text or 'ПОПЪЛНЕТЕ' in p.text or '[' in p.text and ']' in p.text:
        if any(marker in p.text for marker in ['⚠', 'ПОПЪЛНЕТЕ', 'PLACEHOLDER']):
            lines.append(f"  [{i}]: {p.text[:150]}")

with open(OUTPUT, 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))

print(f"Записано в: {OUTPUT}")
print(f"Общо параграфи: {len(doc.paragraphs)}")
print(f"Общо таблици: {len(doc.tables)}")
print(f"Общо редове: {len(lines)}")
