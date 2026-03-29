"""
Скрипт за замяна на placeholder маркерите [⚠️ ПОПЪЛНЕТЕ: ...] в техническото предложение ОП2.
"""

from docx import Document
import re
import os

INPUT = os.path.join(os.path.dirname(__file__), '..', 'GRAFIK DG_KUHNQ',
                     'Кухня_technical_proposal_final_clean (1).docx')
OUTPUT = os.path.join(os.path.dirname(__file__), '..', 'GRAFIK DG_KUHNQ',
                      'Кухня_technical_proposal_final_clean (1).docx')

# Mapping: placeholder text -> replacement text
REPLACEMENTS = {
    '[⚠️ ПОПЪЛНЕТЕ: име на главен проектант]': 'Инж. Георги Писков',
    '[⚠️ ПОПЪЛНЕТЕ: име на складов работник]': 'отговорния складов работник',
    '[⚠️ ПОПЪЛНЕТЕ: име на координатор поддръжка]': 'Инж. Георги Писков',
    '[⚠️ ПОПЪЛНЕТЕ: длъжност на мобилен техник]': 'мобилният сервизен техник на БАРИН АЛП ЕООД',
    '[⚠️ ПОПЪЛНЕТЕ: длъжност на монтажен техник]': 'квалифициран монтажен техник на БАРИН АЛП ЕООД',
    '[⚠️ ПОПЪЛНЕТЕ: име на техник поддръжка]': 'Инж. Георги Писков',
    '[⚠️ ПОПЪЛНЕТЕ: адрес на склад]': 'ул. Миньорска 1',
    '[⚠️ ПОПЪЛНЕТЕ: номер на дежурен телефон]': '0892220535',
}


def replace_in_paragraph(paragraph, replacements):
    """Replace placeholders in a paragraph, preserving run formatting."""
    full_text = paragraph.text
    replaced = False
    
    for old, new in replacements.items():
        if old in full_text:
            replaced = True
            full_text = full_text.replace(old, new)
    
    if replaced:
        # Rebuild runs: keep first run's formatting, put all text there, clear others
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run.text = full_text
            for run in paragraph.runs[1:]:
                run.text = ''
        return True
    return False


def main():
    doc = Document(INPUT)
    total_replaced = 0
    details = []

    # Process body paragraphs
    for i, para in enumerate(doc.paragraphs):
        for old, new in REPLACEMENTS.items():
            if old in para.text:
                replace_in_paragraph(para, {old: new})
                total_replaced += 1
                details.append(f"  Para [{i}]: '{old}' -> '{new}'")

    # Process tables too (just in case)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in REPLACEMENTS.items():
                        if old in para.text:
                            replace_in_paragraph(para, {old: new})
                            total_replaced += 1
                            details.append(f"  Table cell: '{old}' -> '{new}'")

    doc.save(OUTPUT)
    
    # Report
    report = [f"Заменени: {total_replaced} placeholder(s)"]
    report.extend(details)
    
    # Verify no remaining placeholders
    doc2 = Document(OUTPUT)
    remaining = []
    for para in doc2.paragraphs:
        if '⚠️' in para.text or 'ПОПЪЛНЕТЕ' in para.text:
            remaining.append(para.text[:120])
    
    if remaining:
        report.append(f"\n⚠ Останали {len(remaining)} непроменени placeholder(а):")
        for r in remaining:
            report.append(f"  - {r}")
    else:
        report.append("\n✓ Няма останали placeholder маркери.")
    
    output_path = os.path.join(os.path.dirname(__file__), '..', 'GRAFIK DG_KUHNQ', 'tp_fix_report.txt')
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(report))
    
    print('\n'.join(report))


if __name__ == '__main__':
    main()
