# -*- coding: utf-8 -*-
"""
Скрипт за оформление на генериран DOCX файл:
1. Преобразува Markdown заглавия (#, ##, ###, ####) в Word heading стилове
2. Прилага bold/italic форматиране от ** и *
3. Преобразува Markdown таблици в реални Word таблици
4. Премахва [⚠️ ПОПЪЛНЕТЕ: ...] плейсхолдъри, оставяйки само длъжността
5. Задава шрифт Times New Roman 12pt за основен текст
"""

import re
import copy
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

INPUT_FILE = r"e:\VISUAL STUDIO\TEHNICHESKO PREDLOJENIE\ТП_2026-03-26_1637.docx"
OUTPUT_FILE = r"e:\VISUAL STUDIO\TEHNICHESKO PREDLOJENIE\ТП_2026-03-26_1637_formatted.docx"

FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = 12
FONT_SIZE_H1 = 18
FONT_SIZE_H2 = 15
FONT_SIZE_H3 = 13
FONT_SIZE_H4 = 12

# Цветове за заглавия
COLOR_H1 = RGBColor(0x1B, 0x3A, 0x5C)      # тъмносин
COLOR_H2 = RGBColor(0x1F, 0x4E, 0x79)      # среден син
COLOR_H3 = RGBColor(0x2E, 0x2E, 0x2E)      # тъмносив
COLOR_H4 = RGBColor(0x33, 0x33, 0x33)      # сив
COLOR_BODY = RGBColor(0x1A, 0x1A, 0x1A)    # почти черно


def add_bottom_border(paragraph, color="1B3A5C", size="8"):
    """Добавя долна линия (border) под параграф."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr {}><w:bottom w:val="single" w:sz="{}" w:space="1" w:color="{}"/></w:pBdr>'.format(
            nsdecls('w'), size, color))
    pPr.append(pBdr)


def setup_styles(doc):
    """Настройва стиловете на документа."""
    # Основен стил Normal
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(FONT_SIZE_BODY)
    font.color.rgb = COLOR_BODY
    pf = style.paragraph_format
    pf.space_after = Pt(3)
    pf.space_before = Pt(1)
    pf.line_spacing = 1.15

    # Heading стилове с ясна визуална йерархия
    heading_config = [
        (1, FONT_SIZE_H1, True, COLOR_H1, Pt(24), Pt(10)),   # H1: голям, с много място
        (2, FONT_SIZE_H2, True, COLOR_H2, Pt(18), Pt(6)),    # H2: среден
        (3, FONT_SIZE_H3, True, COLOR_H3, Pt(12), Pt(4)),    # H3: по-малък
        (4, FONT_SIZE_H4, True, COLOR_H4, Pt(8), Pt(3)),     # H4: основен размер, bold
    ]
    for level, size, bold, color, space_before, space_after in heading_config:
        style_name = "Heading {}".format(level)
        try:
            hs = doc.styles[style_name]
        except KeyError:
            hs = doc.styles.add_style(style_name, 1)  # WD_STYLE_TYPE.PARAGRAPH
        hf = hs.font
        hf.name = FONT_NAME
        hf.size = Pt(size)
        hf.bold = bold
        hf.color.rgb = color
        hp = hs.paragraph_format
        hp.space_before = space_before
        hp.space_after = space_after
        hp.line_spacing = 1.15


def clean_placeholder(text):
    """
    Премахва [⚠️ ПОПЪЛНЕТЕ: длъжност] и оставя само длъжността.
    Например: '[⚠️ ПОПЪЛНЕТЕ: главен инженер]' -> 'главен инженер'
    """
    # Pattern: [⚠️ ПОПЪЛНЕТЕ: текст] или [⚠️ ПОПЪЛНЕТЕ: текст] 
    pattern = r'\[⚠️\s*ПОПЪЛНЕТЕ:\s*([^\]]*)\]'
    result = re.sub(pattern, lambda m: m.group(1).strip(), text)
    # Също без емоджи
    pattern2 = r'\[\s*ПОПЪЛНЕТЕ[!]*:\s*([^\]]*)\]'
    result = re.sub(pattern2, lambda m: m.group(1).strip(), result)
    # Ако остане "⚠️ ПОПЪЛНЕТЕ: текст" без скоби
    pattern3 = r'⚠️\s*ПОПЪЛНЕТЕ:\s*'
    result = re.sub(pattern3, '', result)
    return result


def determine_heading_level(text):
    """Определя нивото на heading от markdown синтаксис."""
    # Премахваме водещи ** маркери ако има
    stripped = text.strip()
    # Премахваме обграждащи **
    temp = stripped
    if temp.startswith("**"):
        temp = temp[2:]
    if temp.endswith("**"):
        temp = temp[:-2]
    temp = temp.strip()
    
    # Преброяваме ВСИЧКИ водещи # символи
    hash_count = 0
    check = temp
    while check.startswith("#"):
        hash_count += 1
        check = check[1:]
    
    if hash_count == 0:
        # Проверяваме оригиналния текст
        check2 = stripped
        while check2.startswith("#"):
            hash_count += 1
            check2 = check2[1:]
        if hash_count == 0:
            return 0, text
        clean_text = check2.lstrip()
    else:
        clean_text = check.lstrip()
    
    # Ограничаваме до ниво 4 (Word поддържа до Heading 4 стандартно)
    level = min(hash_count, 4)
    
    return level, clean_text


def apply_inline_formatting(paragraph, text, base_bold=False, base_italic=False, font_name=FONT_NAME, font_size=None):
    """
    Парсва markdown inline форматиране (**bold**, *italic*, ***bold italic***) 
    и добавя runs с правилното форматиране.
    """
    # Премахваме placeholder-и
    text = clean_placeholder(text)

    # Стратегия: първо обработваме *** (bold+italic), после ** (bold), после * (italic)
    # Използваме token-based подход за по-добра надеждност
    
    # Парсваме с приоритет: ***text*** > **text** > *text*
    parts = []
    pos = 0
    
    # Комбиниран pattern с приоритет
    pattern = r'\*{3}(.+?)\*{3}|\*{2}(.+?)\*{2}|\*([^*]+?)\*'

    for match in re.finditer(pattern, text):
        # Текст преди мача
        if match.start() > pos:
            parts.append(('normal', text[pos:match.start()]))
        
        if match.group(1) is not None:
            parts.append(('bold_italic', match.group(1)))
        elif match.group(2) is not None:
            parts.append(('bold', match.group(2)))
        elif match.group(3) is not None:
            parts.append(('italic', match.group(3)))
        pos = match.end()

    # Остатък
    if pos < len(text):
        parts.append(('normal', text[pos:]))

    # Ако няма маркери
    if not parts:
        parts = [('normal', text)]

    # Почистваме стряй * символи от normal части
    cleaned_parts = []
    for fmt, content in parts:
        if fmt == 'normal':
            # Премахваме самостоятелни * (не част от **)
            content = content.replace("***", "").replace("**", "")
            # Премахваме водещи и завършващи * само ако са stray
            content = content.strip('*')
            # Премахваме * ако е единственият символ
            if content.strip() == '*':
                continue
        if not content or not content.strip():
            continue
        cleaned_parts.append((fmt, content))
    
    if not cleaned_parts:
        cleaned_parts = [('normal', '')]

    for fmt, content in cleaned_parts:
        if not content:
            continue
        run = paragraph.add_run(content)
        run.font.name = font_name
        if font_size:
            run.font.size = Pt(font_size)
        run.font.color.rgb = COLOR_BODY

        if fmt == 'bold' or base_bold:
            run.bold = True
        if fmt == 'italic' or base_italic:
            run.italic = True
        if fmt == 'bold_italic':
            run.bold = True
            run.italic = True


def parse_markdown_table(paragraphs, start_idx):
    """
    Парсва markdown таблица от списък параграфи.
    Връща (rows_data, end_idx) където rows_data е списък от списъци с клетки.
    """
    rows = []
    idx = start_idx
    while idx < len(paragraphs):
        text = paragraphs[idx].text.strip()
        if not text.startswith("|") and "---|" not in text:
            break
        # Пропускаме separator редове (|---|---|)
        if re.match(r'^\|[\s\-:|]+\|$', text):
            idx += 1
            continue
        # Парсваме клетки
        cells = [c.strip() for c in text.split("|")]
        # Премахваме празни елементи от началото и края
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        if cells:
            rows.append(cells)
        idx += 1
    return rows, idx


def create_word_table(doc, rows_data, insert_before_element):
    """Създава Word таблица с форматиране."""
    if not rows_data or not rows_data[0]:
        return None

    num_cols = max(len(row) for row in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Стил на таблицата
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = table.cell(i, j)
            # Изчистваме default параграф
            cell.text = ''
            p = cell.paragraphs[0]

            # Премахваме placeholder-и
            cell_text = clean_placeholder(cell_text)
            
            # Заглавен ред (първи ред) - bold
            is_header = (i == 0)
            
            # Премахваме ** маркери от текста
            clean_text = cell_text.replace("**", "")
            
            run = p.add_run(clean_text)
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            if is_header:
                run.bold = True
                # Оцветяване на header ред
                shading = parse_xml('<w:shd {} w:fill="2E74B5" w:val="clear"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading)
                run.font.color.rgb = RGBColor(255, 255, 255)
            
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

    # Автоматична ширина
    table.autofit = True
    
    return table


def is_bullet_line(text):
    """Проверява дали е маркиран списък."""
    stripped = text.lstrip()
    return stripped.startswith("- ") or stripped.startswith("• ") or re.match(r'^\d+\.\s', stripped)


def get_bullet_text(text):
    """Извлича текста от маркиран списък."""
    stripped = text.lstrip()
    if stripped.startswith("- "):
        return stripped[2:]
    if stripped.startswith("• "):
        return stripped[2:]
    m = re.match(r'^\d+\.\s', stripped)
    if m:
        return stripped[m.end():]
    return text


def main():
    print("Зареждане на документа...")
    doc = Document(INPUT_FILE)
    
    print("Настройване на стилове...")
    setup_styles(doc)
    
    # Събираме всички параграфи с техния текст и позиция
    paragraphs = doc.paragraphs
    total = len(paragraphs)
    print("Общо параграфи: {}".format(total))
    
    # Първо идентифицираме markdown таблици
    # За да ги заменим, трябва да работим на ниво XML
    
    body = doc.element.body
    
    # Събираме всички елементи в body
    elements = list(body)
    
    # Създаваме нов документ
    new_doc = Document()
    setup_styles(new_doc)
    
    # Задаваме margins
    for section in new_doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    # Изтриваме default празен параграф
    if new_doc.paragraphs:
        p = new_doc.paragraphs[0]
        p_element = p._element
        p_element.getparent().remove(p_element)
    
    i = 0
    table_count = 0
    heading_count = 0
    placeholder_count = 0
    
    while i < total:
        para = paragraphs[i]
        text = para.text
        stripped = text.strip()
        
        # Проверяваме за празен параграф
        if not stripped:
            i += 1
            continue
        
        # Проверяваме за начало на markdown таблица
        if stripped.startswith("|") or "---|" in stripped:
            # Парсваме таблицата
            rows_data, end_idx = parse_markdown_table(paragraphs, i)
            if rows_data and len(rows_data) > 0:
                table_count += 1
                # Добавяме малко пространство
                spacer = new_doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(4)
                spacer.paragraph_format.space_after = Pt(0)
                
                # Създаваме таблица в новия документ
                num_cols = max(len(row) for row in rows_data)
                table = new_doc.add_table(rows=len(rows_data), cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                try:
                    table.style = 'Table Grid'
                except:
                    pass
                
                for ri, row_data in enumerate(rows_data):
                    for ci, cell_text in enumerate(row_data):
                        if ci >= num_cols:
                            break
                        cell = table.cell(ri, ci)
                        cell.text = ''
                        p = cell.paragraphs[0]
                        
                        cell_text = clean_placeholder(cell_text)
                        clean_text = cell_text.replace("**", "")
                        
                        is_header = (ri == 0)
                        
                        run = p.add_run(clean_text)
                        run.font.name = FONT_NAME
                        run.font.size = Pt(10)
                        
                        if is_header:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            shading = parse_xml('<w:shd {} w:fill="2E74B5" w:val="clear"/>'.format(nsdecls('w')))
                            cell._tc.get_or_add_tcPr().append(shading)
                        else:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            # Алтернативен цвят на редовете
                            if ri % 2 == 0:
                                shading = parse_xml('<w:shd {} w:fill="D6E4F0" w:val="clear"/>'.format(nsdecls('w')))
                                cell._tc.get_or_add_tcPr().append(shading)
                        
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.space_before = Pt(1)
                        p.paragraph_format.space_after = Pt(1)
                
                table.autofit = True
                
                # Добавяме малко пространство след таблицата
                spacer2 = new_doc.add_paragraph()
                spacer2.paragraph_format.space_before = Pt(0)
                spacer2.paragraph_format.space_after = Pt(4)
                
                i = end_idx
                continue
        
        # Проверяваме за placeholder
        if "ПОПЪЛНЕТЕ" in text:
            placeholder_count += 1
        
        # Проверяваме за heading
        heading_level, heading_text = determine_heading_level(stripped)
        
        if heading_level > 0:
            heading_count += 1
            heading_text = clean_placeholder(heading_text)
            # Премахваме markdown bold/italic маркери от заглавията
            heading_text = heading_text.replace("***", "").replace("**", "").replace("*", "")
            
            p = new_doc.add_paragraph(style="Heading {}".format(heading_level))
            
            sizes = {1: FONT_SIZE_H1, 2: FONT_SIZE_H2, 3: FONT_SIZE_H3, 4: FONT_SIZE_H4}
            colors = {1: COLOR_H1, 2: COLOR_H2, 3: COLOR_H3, 4: COLOR_H4}
            
            # H1: ГЛАВНО ЗАГЛАВИЕ — центрирано, с долна линия
            if heading_level == 1:
                run = p.add_run(heading_text.upper())
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_H1)
                run.bold = True
                run.font.color.rgb = COLOR_H1
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(30)
                p.paragraph_format.space_after = Pt(12)
                add_bottom_border(p, "1B3A5C", "12")
            
            # H2: Основни секции — с долна линия
            elif heading_level == 2:
                run = p.add_run(heading_text)
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_H2)
                run.bold = True
                run.font.color.rgb = COLOR_H2
                p.paragraph_format.space_before = Pt(22)
                p.paragraph_format.space_after = Pt(8)
                add_bottom_border(p, "1F4E79", "6")
            
            # H3: Подсекции — леко отместени
            elif heading_level == 3:
                run = p.add_run(heading_text)
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_H3)
                run.bold = True
                run.font.color.rgb = COLOR_H3
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Cm(0.5)
            
            # H4: Под-подсекции — italic + отместване
            elif heading_level == 4:
                run = p.add_run(heading_text)
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_H4)
                run.bold = True
                run.font.color.rgb = COLOR_H4
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Cm(1.0)
            
            i += 1
            continue
        
        # Bullet list
        if is_bullet_line(stripped):
            bullet_text = get_bullet_text(stripped)
            p = new_doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, bullet_text, font_size=FONT_SIZE_BODY)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(1.5)
            i += 1
            continue
        
        # Обикновен параграф
        p = new_doc.add_paragraph()
        apply_inline_formatting(p, stripped, font_size=FONT_SIZE_BODY)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.first_line_indent = Cm(0.75)
        
        i += 1
    
    print("Преобразувани заглавия: {}".format(heading_count))
    print("Създадени таблици: {}".format(table_count))
    print("Коригирани placeholder-и: {}".format(placeholder_count))
    
    # Запазваме
    print("Запазване на {}...".format(OUTPUT_FILE))
    new_doc.save(OUTPUT_FILE)
    print("Готово!")


if __name__ == "__main__":
    main()
