# -*- coding: utf-8 -*-
"""
Скрипт за оформление на final_clean_file6.docx:
- Задава Times New Roman шрифтове с правилни размери
- Форматира Heading 1/2 стилове с цветове, линии, разстояния
- Форматира bold Normal параграфи като визуални под-заглавия
- Премахва ПОПЪЛНЕТЕ плейсхолдъри
- Форматира bullets и numbered lists
- Оцветява съществуващи таблици
- Добавява номера на страниците в footer
- Разделя секции с разстояние (без page breaks)
"""

import re
import sys
import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

INPUT_FILE = r"e:\VISUAL STUDIO\TEHNICHESKO PREDLOJENIE\final_clean_file6.docx"
OUTPUT_FILE = r"e:\VISUAL STUDIO\TEHNICHESKO PREDLOJENIE\final_clean_file6_formatted.docx"

FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = 12
FONT_SIZE_H1 = 16
FONT_SIZE_H2 = 14
FONT_SIZE_H3 = 13
FONT_SIZE_H4 = 12
FONT_SIZE_SUBHEADING = 11
FONT_SIZE_CATEGORY = 10

COLOR_H1 = RGBColor(0x1B, 0x3A, 0x5C)
COLOR_H2 = RGBColor(0x1F, 0x4E, 0x79)
COLOR_BODY = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_SUBHEADING = RGBColor(0x1F, 0x4E, 0x79)
COLOR_LABEL = RGBColor(0x2E, 0x4E, 0x6E)

RE_PLACEHOLDER = re.compile(r'\[⚠️\s*ПОПЪЛНЕТЕ:\s*([^\]]*)\]')
RE_PLACEHOLDER2 = re.compile(r'\[\s*ПОПЪЛНЕТЕ[!]*:\s*([^\]]*)\]')
RE_PLACEHOLDER3 = re.compile(r'⚠️\s*ПОПЪЛНЕТЕ:\s*')
RE_NUMBERED = re.compile(r'^(\d+)\.\s')
RE_DEYNOST = re.compile(r'^Дейност\s+\d+', re.IGNORECASE)
RE_ETAP_SUB = re.compile(r'^(Етап|Стъпка|Фаза|Подетап)\s+\d+', re.IGNORECASE)


def add_bottom_border(paragraph, color="1B3A5C", size="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr {}><w:bottom w:val="single" w:sz="{}" w:space="1" w:color="{}"/></w:pBdr>'.format(
            nsdecls('w'), size, color))
    pPr.append(pBdr)


def add_left_border(paragraph, color="2E74B5", size="18"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr {}><w:left w:val="single" w:sz="{}" w:space="4" w:color="{}"/></w:pBdr>'.format(
            nsdecls('w'), size, color))
    pPr.append(pBdr)


def add_paragraph_shading(paragraph, color="EDF2F9"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = parse_xml('<w:shd {} w:fill="{}" w:val="clear"/>'.format(nsdecls('w'), color))
    pPr.append(shd)


def clean_placeholder(text):
    result = RE_PLACEHOLDER.sub(lambda m: m.group(1).strip(), text)
    result = RE_PLACEHOLDER2.sub(lambda m: m.group(1).strip(), result)
    result = RE_PLACEHOLDER3.sub('', result)
    return result


def set_run_font(run, size=None, bold=None, italic=None, color=None):
    """Задава шрифт на run."""
    run.font.name = FONT_NAME
    # Force font for East Asian and Complex Script
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml('<w:rFonts {} w:ascii="{}" w:hAnsi="{}" w:cs="{}" w:eastAsia="{}"/>'.format(
            nsdecls('w'), FONT_NAME, FONT_NAME, FONT_NAME, FONT_NAME))
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)
        rFonts.set(qn('w:cs'), FONT_NAME)
        rFonts.set(qn('w:eastAsia'), FONT_NAME)
    
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_page_number_footer(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        # Clear existing
        for run in p.runs:
            run.text = ''
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run1 = p.add_run("— ")
        set_run_font(run1, size=9, color=RGBColor(0x88, 0x88, 0x88))
        
        fld = parse_xml(
            '<w:fldSimple {} w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>'.format(nsdecls('w')))
        p._p.append(fld)
        
        run2 = p.add_run(" / ")
        set_run_font(run2, size=9, color=RGBColor(0x88, 0x88, 0x88))
        
        fld2 = parse_xml(
            '<w:fldSimple {} w:instr=" NUMPAGES "><w:r><w:t>1</w:t></w:r></w:fldSimple>'.format(nsdecls('w')))
        p._p.append(fld2)
        
        run3 = p.add_run(" —")
        set_run_font(run3, size=9, color=RGBColor(0x88, 0x88, 0x88))


def is_bold_paragraph(para):
    """Проверява дали параграф е изцяло bold (всички runs с текст са bold)."""
    runs_with_text = [r for r in para.runs if r.text.strip()]
    if not runs_with_text:
        return False
    return all(r.bold for r in runs_with_text)


def is_deynost_heading(text):
    """Проверява дали е заглавие тип 'Дейност N:...'"""
    return bool(RE_DEYNOST.match(text.strip()))


def is_label_heading(text):
    """Проверява дали е кратко bold заглавие-етикет."""
    t = text.strip()
    label_patterns = [
        'Конкретни действия', 'Необходимо оборудване', 'Отговорен експерт',
        'Технологична последователност', 'Необходими машини', 'Методи за',
        'Управление на отпадъци', 'Техники за', 'Специализирано оборудване',
        'Временно съхранение', 'Мерки за', 'Времеви план', 'Критерии за',
        'Процедури за', 'Система за', 'Контролни точки', 'Резултати от',
        'Условия за', 'Предпазни мерки', 'Документация', 'Ресурсно осигуряване',
        'Изисквания към', 'Организация на', 'Защита на', 'Координация',
        'Комуникация', 'Конкретни мерки', 'Основни принципи', 'Последователност',
        'Материали и инструменти', 'Ключови', 'Превантивни мерки',
    ]
    for pattern in label_patterns:
        if t.startswith(pattern):
            return True
    # Кратък bold текст (2-6 думи) вероятно е етикет
    if t.endswith(':') and len(t.split()) <= 8:
        return True
    return False


def format_table(table):
    """Форматира съществуваща Word таблица с хубаво оцветяване."""
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

            if ri == 0:
                # Header row
                shading = parse_xml('<w:shd {} w:fill="2E74B5" w:val="clear"/>'.format(nsdecls('w')))
                # Remove existing shading first
                tcPr = cell._tc.get_or_add_tcPr()
                existing_shd = tcPr.find(qn('w:shd'))
                if existing_shd is not None:
                    tcPr.remove(existing_shd)
                tcPr.append(shading)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True
            else:
                # Zebra striping
                if ri % 2 == 0:
                    shading = parse_xml('<w:shd {} w:fill="E8EFF7" w:val="clear"/>'.format(nsdecls('w')))
                    tcPr = cell._tc.get_or_add_tcPr()
                    existing_shd = tcPr.find(qn('w:shd'))
                    if existing_shd is not None:
                        tcPr.remove(existing_shd)
                    tcPr.append(shading)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)


def main():
    print("Зареждане на документа...")
    doc = Document(INPUT_FILE)
    
    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    
    # Page numbers
    add_page_number_footer(doc)
    
    # Настройваме стиловете
    print("Настройване на стилове...")
    
    # Normal style
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE_BODY)
    style.font.color.rgb = COLOR_BODY
    pf = style.paragraph_format
    pf.space_after = Pt(3)
    pf.space_before = Pt(1)
    pf.line_spacing = 1.15
    
    # Heading styles
    for level, size, color, sb, sa in [
        (1, FONT_SIZE_H1, COLOR_H1, Pt(28), Pt(10)),
        (2, FONT_SIZE_H2, COLOR_H2, Pt(20), Pt(8)),
    ]:
        style_name = "Heading {}".format(level)
        try:
            hs = doc.styles[style_name]
            hs.font.name = FONT_NAME
            hs.font.size = Pt(size)
            hs.font.bold = True
            hs.font.color.rgb = color
            hs.paragraph_format.space_before = sb
            hs.paragraph_format.space_after = sa
            hs.paragraph_format.line_spacing = 1.15
        except KeyError:
            pass
    
    # Форматиране на съществуващи таблици
    print("Форматиране на {} таблици...".format(len(doc.tables)))
    for table in doc.tables:
        format_table(table)
    
    # Обработка на параграфи
    paras = doc.paragraphs
    total = len(paras)
    print("Обработка на {} параграфи...".format(total))
    
    placeholder_count = 0
    subheading_count = 0
    deynost_count = 0
    label_count = 0
    heading_formatted = 0
    bullet_count = 0
    spacer_count = 0
    md_heading_count = 0
    
    prev_style = None
    
    for idx, para in enumerate(paras):
        text = para.text.strip()
        if not text:
            continue
        
        style_name = para.style.name if para.style else 'Normal'
        
        # ===== CLEAN PLACEHOLDERS =====
        if "ПОПЪЛНЕТЕ" in para.text:
            placeholder_count += 1
            # Текстът е разделен между runs — събираме целия текст, чистим, записваме обратно
            full_text = para.text
            cleaned = clean_placeholder(full_text)
            if full_text != cleaned:
                # Запазваме форматирането на първия run, изтриваме останалите
                if para.runs:
                    first_run = para.runs[0]
                    first_run.text = cleaned
                    for run in para.runs[1:]:
                        run.text = ''
        
        # ===== MARKDOWN HEADING (#) =====
        if text.startswith('#') and style_name == 'Normal':
            md_heading_count += 1
            # Count hash marks
            hash_count = 0
            temp = text
            while temp.startswith('#'):
                hash_count += 1
                temp = temp[1:]
            clean_text = temp.strip()
            
            # Set heading style
            level = min(hash_count, 2)
            try:
                para.style = doc.styles["Heading {}".format(level)]
            except:
                pass
            
            # Replace text
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = clean_text.upper() if level == 1 else clean_text
            else:
                run = para.add_run(clean_text.upper() if level == 1 else clean_text)
            
            # Format
            for run in para.runs:
                if run.text.strip():
                    set_run_font(run, size=FONT_SIZE_H1 if level == 1 else FONT_SIZE_H2,
                                bold=True, color=COLOR_H1 if level == 1 else COLOR_H2)
            
            if level == 1:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(28)
                para.paragraph_format.space_after = Pt(12)
                add_bottom_border(para, "1B3A5C", "12")
            else:
                para.paragraph_format.space_before = Pt(22)
                para.paragraph_format.space_after = Pt(8)
                add_bottom_border(para, "1F4E79", "6")
            
            prev_style = 'heading'
            continue
        
        # ===== HEADING 1 =====
        if style_name == 'Heading 1':
            heading_formatted += 1
            for run in para.runs:
                set_run_font(run, size=FONT_SIZE_H1, bold=True, color=COLOR_H1)
            
            para.paragraph_format.space_before = Pt(28)
            para.paragraph_format.space_after = Pt(10)
            add_bottom_border(para, "1B3A5C", "10")
            
            # Разстояние преди H1 ако предишният не е heading
            if prev_style not in ('heading', None):
                para.paragraph_format.space_before = Pt(36)
            
            prev_style = 'heading'
            continue
        
        # ===== HEADING 2 =====
        if style_name == 'Heading 2':
            heading_formatted += 1
            for run in para.runs:
                set_run_font(run, size=FONT_SIZE_H2, bold=True, color=COLOR_H2)
            
            para.paragraph_format.space_before = Pt(20)
            para.paragraph_format.space_after = Pt(8)
            add_bottom_border(para, "1F4E79", "6")
            
            prev_style = 'heading'
            continue
        
        # ===== BOLD PARAGRAPH (sub-heading) =====
        if style_name == 'Normal' and is_bold_paragraph(para):
            clean_text = text
            
            if is_deynost_heading(clean_text):
                # Дейност N: ... — виден под-заглавие с фон и лява линия
                deynost_count += 1
                add_paragraph_shading(para, "EDF2F9")
                add_left_border(para, "2E74B5", "18")
                para.paragraph_format.space_before = Pt(14)
                para.paragraph_format.space_after = Pt(4)
                para.paragraph_format.left_indent = Cm(0.3)
                for run in para.runs:
                    set_run_font(run, size=FONT_SIZE_SUBHEADING, bold=True, color=COLOR_SUBHEADING)
                prev_style = 'subheading'
            
            elif is_label_heading(clean_text):
                # Етикет: "Конкретни действия:", "Необходимо оборудване:" и т.н.
                label_count += 1
                add_left_border(para, "95B3D7", "12")
                para.paragraph_format.space_before = Pt(8)
                para.paragraph_format.space_after = Pt(3)
                para.paragraph_format.left_indent = Cm(0.5)
                for run in para.runs:
                    set_run_font(run, size=FONT_SIZE_SUBHEADING, bold=True, 
                                color=RGBColor(0x3A, 0x5A, 0x8C))
                prev_style = 'label'
            
            else:
                # Друг bold параграф — generic sub-heading
                subheading_count += 1
                add_paragraph_shading(para, "EDF2F9")
                para.paragraph_format.space_before = Pt(10)
                para.paragraph_format.space_after = Pt(3)
                para.paragraph_format.left_indent = Cm(0.3)
                for run in para.runs:
                    set_run_font(run, size=FONT_SIZE_SUBHEADING, bold=True, color=COLOR_SUBHEADING)
                prev_style = 'subheading'
            
            continue
        
        # ===== BULLET LIST =====
        if text.startswith('- ') or text.startswith('• '):
            bullet_count += 1
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            if para.paragraph_format.left_indent is None or para.paragraph_format.left_indent == 0:
                para.paragraph_format.left_indent = Cm(1.5)
            for run in para.runs:
                set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BODY)
            prev_style = 'bullet'
            continue
        
        # ===== NUMBERED LIST =====
        if RE_NUMBERED.match(text):
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            if para.paragraph_format.left_indent is None or para.paragraph_format.left_indent == 0:
                para.paragraph_format.left_indent = Cm(1.0)
            for run in para.runs:
                set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BODY)
            prev_style = 'numbered'
            continue
        
        # ===== NORMAL PARAGRAPH =====
        for run in para.runs:
            set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BODY)
        
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(3)
        
        # Отстъп на първи ред за обикновен текст (не след heading/subheading)
        if prev_style not in ('heading', 'subheading', 'label'):
            if para.paragraph_format.first_line_indent is None:
                para.paragraph_format.first_line_indent = Cm(0.75)
        
        prev_style = 'normal'
    
    print("=" * 50)
    print("Heading-и форматирани: {}".format(heading_formatted))
    print("Markdown заглавия (#): {}".format(md_heading_count))
    print("Дейност заглавия: {}".format(deynost_count))
    print("Label заглавия: {}".format(label_count))
    print("Други bold под-заглавия: {}".format(subheading_count))
    print("Bullets: {}".format(bullet_count))
    print("Коригирани placeholder-и: {}".format(placeholder_count))
    print("Таблици форматирани: {}".format(len(doc.tables)))
    print("=" * 50)
    
    print("Запазване на {}...".format(OUTPUT_FILE))
    doc.save(OUTPUT_FILE)
    print("Готово!")


if __name__ == "__main__":
    main()
