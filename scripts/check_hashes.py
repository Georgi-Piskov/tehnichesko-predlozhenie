import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

doc = Document(r"e:\VISUAL STUDIO\TEHNICHESKO PREDLOJENIE\ТП_2026-03-26_1637.docx")

print("=== ПАРАГРАФИ С # МАРКЕРИ (подробно) ===")
count = 0
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    # Проверяваме дали след премахване на ** все още има # в началото
    temp = t
    if temp.startswith("**"):
        temp = temp[2:]
    temp = temp.strip()
    if temp.startswith("#") and not t.startswith("#"):
        count += 1
        if count <= 20:
            print("{}: '{}'".format(i, t[:200]))
print("Total # inside ** markers:", count)

print("\n=== HEADING-LIKE LINES NOT STARTING WITH # ===")
count2 = 0
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t.startswith("#") and t.startswith("**") and ("# " in t or "## " in t):
        count2 += 1
        if count2 <= 10:
            print("{}: '{}'".format(i, t[:200]))
print("Total:", count2)
