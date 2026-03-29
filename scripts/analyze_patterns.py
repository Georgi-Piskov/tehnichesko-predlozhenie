import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

doc = Document(r"e:\VISUAL STUDIO\TEHNICHESKO PREDLOJENIE\ТП_2026-03-26_1637.docx")

# Показваме разнообразие от параграфи за да видим какви label patterns има
print("=== ПРИМЕРНИ ПАРАГРАФИ (200-350) ===")
for i in range(200, min(350, len(doc.paragraphs))):
    t = doc.paragraphs[i].text.strip()
    if t:
        print("{}: {}".format(i, t[:200]))

print("\n\n=== ПРИМЕРНИ ПАРАГРАФИ (500-650) ===")
for i in range(500, min(650, len(doc.paragraphs))):
    t = doc.paragraphs[i].text.strip()
    if t:
        print("{}: {}".format(i, t[:200]))

print("\n\n=== LABEL PATTERNS (Дейност, Контрол, KPI, etc.) ===")
labels = {}
import re
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    # Find patterns like "- **Label:** text" or "**Label:** text"
    m = re.match(r'^[-•]?\s*\*?\*?\s*(\w[\w\s\-]{2,25}?):\*?\*?\s', t)
    if m:
        label = m.group(1).strip()
        if label not in labels:
            labels[label] = []
        if len(labels[label]) < 2:
            labels[label].append("{}: {}".format(i, t[:120]))

print("Unique labels found:", len(labels))
for label, examples in sorted(labels.items()):
    print("\n  Label: '{}'".format(label))
    for ex in examples:
        print("    {}".format(ex))

print("\n\n=== BOLD-ONLY PARAGRAPHS (entire line is **text**) ===")
bold_lines = 0
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith("**") and t.endswith("**") and t.count("**") == 2:
        bold_lines += 1
        if bold_lines <= 15:
            print("{}: {}".format(i, t[:150]))
print("Total bold-only lines:", bold_lines)
