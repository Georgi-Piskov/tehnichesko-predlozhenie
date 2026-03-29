import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

doc = Document(r"e:\VISUAL STUDIO\TEHNICHESKO PREDLOJENIE\ТП_2026-03-26_1637_formatted.docx")

# Проверяваме за параграфи с mix от bold и non-bold runs
print("=== ПАРАГРАФИ С MIXED FORMATTING (first 10) ===")
count = 0
for p in doc.paragraphs:
    if p.style.name == 'Normal' and len(p.runs) > 1:
        has_bold = any(r.bold for r in p.runs)
        has_nonbold = any(not r.bold for r in p.runs)
        if has_bold and has_nonbold:
            count += 1
            if count <= 10:
                info = []
                for r in p.runs[:5]:
                    info.append("[bold={}: '{}']".format(r.bold, r.text[:30]))
                print("  {}".format(" | ".join(info)))
print("Total mixed:", count)

print("\n=== ПАРАГРАФИ С НЕ-BOLD ТЕКСТ (first 5) ===")
count2 = 0  
for p in doc.paragraphs:
    if p.style.name in ('Normal', 'List Bullet'):
        for r in p.runs:
            if not r.bold and r.text.strip():
                count2 += 1
                if count2 <= 5:
                    print("  стил={}: '{}'".format(p.style.name, r.text[:80]))
                break
print("Total non-bold paragraphs:", count2)

print("\n=== MARKDOWN ** MARKERS REMAINING ===")
star_count = 0
for p in doc.paragraphs:
    if "**" in p.text:
        star_count += 1
        if star_count <= 5:
            print("  {}".format(p.text[:120]))
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            if "**" in cell.text:
                star_count += 1
print("Total with ** remaining:", star_count)

print("\n=== MARKDOWN # MARKERS REMAINING ===")
hash_count = 0
for p in doc.paragraphs:
    if p.text.strip().startswith("#"):
        hash_count += 1
        if hash_count <= 3:
            print("  {}".format(p.text[:120]))
print("Total with # remaining:", hash_count)
