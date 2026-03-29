import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

doc = Document(r"e:\VISUAL STUDIO\TEHNICHESKO PREDLOJENIE\ТП_2026-03-26_1637_formatted.docx")

print("=== LINES STARTING WITH # (style and text) ===")
count = 0
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith("#"):
        count += 1
        if count <= 30:
            print("{}: style='{}' text='{}'".format(i, p.style.name, t[:150]))
print("Total:", count)

# Are these headings or normal?
heading_hash = 0
normal_hash = 0
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith("#"):
        if 'Heading' in p.style.name:
            heading_hash += 1
        else:
            normal_hash += 1
print("\nWith Heading style:", heading_hash)
print("With Normal style:", normal_hash)
