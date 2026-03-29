import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

doc = Document(r"e:\VISUAL STUDIO\TEHNICHESKO PREDLOJENIE\ТП_2026-03-26_1637_formatted.docx")

print("=== СТИЛОВЕ В УПОТРЕБА ===")
styles_used = {}
for p in doc.paragraphs:
    name = p.style.name
    styles_used[name] = styles_used.get(name, 0) + 1
for name, count in sorted(styles_used.items()):
    print("  {}: {} параграфа".format(name, count))

print("\n=== ТАБЛИЦИ ===")
print("Общо таблици:", len(doc.tables))
for i, t in enumerate(doc.tables[:5]):
    print("Таблица {}: {} реда x {} колони".format(i, len(t.rows), len(t.columns)))
    if t.rows:
        header = [c.text[:30] for c in t.rows[0].cells]
        print("  Header:", header)

print("\n=== ПРОВЕРКА ЗА ПОПЪЛНЕТЕ ===")
pop_count = 0
for p in doc.paragraphs:
    if "ПОПЪЛНЕТЕ" in p.text:
        pop_count += 1
        print("  ОСТАВА:", p.text[:150])
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            if "ПОПЪЛНЕТЕ" in cell.text:
                pop_count += 1
                if pop_count <= 10:
                    print("  ОСТАВА (таблица):", cell.text[:150])
print("Общо останали ПОПЪЛНЕТЕ:", pop_count)

print("\n=== ПРИМЕР ЗАГЛАВИЯ (първи 10) ===")
count = 0
for p in doc.paragraphs:
    if 'Heading' in p.style.name:
        count += 1
        if count <= 10:
            print("  {}: {}".format(p.style.name, p.text[:100]))

print("\n=== ПРОВЕРКА ШРИФТОВЕ ===")
for p in doc.paragraphs[:20]:
    for r in p.runs[:1]:
        print("  стил={}, font={}, size={}, bold={}".format(p.style.name, r.font.name, r.font.size, r.bold))
        break

print("\n=== ОБЩО ПАРАГРАФИ:", len(doc.paragraphs))
