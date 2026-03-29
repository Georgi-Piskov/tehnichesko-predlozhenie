import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

doc = Document(r"e:\VISUAL STUDIO\TEHNICHESKO PREDLOJENIE\ТП_2026-03-26_1637_formatted.docx")

# Check for stray * runs
print("=== СТИЛОВЕ ===")
styles_used = {}
for p in doc.paragraphs:
    name = p.style.name
    styles_used[name] = styles_used.get(name, 0) + 1
for name, count in sorted(styles_used.items()):
    print("  {}: {}".format(name, count))

print("\n=== STRAY * CHARACTERS ===")
stray_count = 0
for p in doc.paragraphs:
    for r in p.runs:
        if r.text.strip() == '*':
            stray_count += 1
print("Total stray * runs:", stray_count)

print("\n=== REMAINING ПОПЪЛНЕТЕ ===")
pop = 0
for p in doc.paragraphs:
    if "ПОПЪЛНЕТЕ" in p.text:
        pop += 1
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            if "ПОПЪЛНЕТЕ" in cell.text:
                pop += 1
print("Total:", pop)

print("\n=== REMAINING ** MARKERS ===")
stars = 0
for p in doc.paragraphs:
    if "**" in p.text:
        stars += 1
print("Total:", stars)

print("\n=== REMAINING # MARKERS ===")
hashes = 0
for p in doc.paragraphs:
    if p.text.strip().startswith("#"):
        hashes += 1
print("Total:", hashes)

print("\n=== STATS ===")
print("Total paragraphs:", len(doc.paragraphs))
print("Total tables:", len(doc.tables))
print("Headings:", sum(v for k,v in styles_used.items() if 'Heading' in k))
print("Bullets:", styles_used.get('List Bullet', 0))
print("Normal:", styles_used.get('Normal', 0))
