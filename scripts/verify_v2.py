# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

doc = Document(r'e:\VISUAL STUDIO\TEHNICHESKO PREDLOJENIE\ТП_2026-03-26_1637_formatted_v2.docx')

paras = doc.paragraphs
tables = doc.tables
print("Параграфи: {}".format(len(paras)))
print("Таблици: {}".format(len(tables)))

stray_star = sum(1 for p in paras if '*' in p.text)
populate = sum(1 for p in paras if 'ПОПЪЛНЕТЕ' in p.text)
hash_marks = sum(1 for p in paras if '#' in p.text)
print("Параграфи с *: {}".format(stray_star))
print("ПОПЪЛНЕТЕ: {}".format(populate))
print("Параграфи с #: {}".format(hash_marks))

# Стилове breakdown
style_counts = {}
for p in paras:
    s = p.style.name if p.style else 'None'
    style_counts[s] = style_counts.get(s, 0) + 1
print("\nСтилове:")
for s, c in sorted(style_counts.items(), key=lambda x: -x[1]):
    print("  {}: {}".format(s, c))

# Sample paragraphs
print("\n=== ПРИМЕРНИ ПАРАГРАФИ (първи 25 непразни) ===")
count = 0
for p in paras:
    if not p.text.strip():
        continue
    count += 1
    if count > 25:
        break
    style = p.style.name if p.style else 'None'
    has_border = 'pBdr' in p._p.xml
    has_shading = 'shd' in p._p.xml
    flags = ""
    if has_border: flags += "[border]"
    if has_shading: flags += "[shaded]"
    runs_detail = []
    for r in p.runs:
        b = "B" if r.bold else ""
        it = "I" if r.italic else ""
        sz = str(r.font.size.pt) if r.font.size else "?"
        runs_detail.append("{}{}{}pt".format(b, it, sz))
    print("{}. [{}] {} {}".format(count, style, flags, p.text[:90]))
    print("   Runs: {}".format(", ".join(runs_detail[:5])))
